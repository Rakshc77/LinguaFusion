from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Dict, Tuple

from docx import Document

from backend.services.file_reader_service import read_plain_text, extract_text_from_file, extract_csv_text
from backend.services.language_service import detect_text_language
from backend.services.translation_service import translate_with_views, normalize_lang


SUPPORTED_FORMAT_PRESERVE_EXTENSIONS = {".txt", ".md", ".csv", ".docx", ".pdf"}




def _font_candidates() -> list[Path]:
    candidates = [
        Path(r"C:\Windows\Fonts\Nirmala.ttf"),
        Path(r"C:\Windows\Fonts\NirmalaB.ttf"),
        Path(r"C:\Windows\Fonts\Mangal.ttf"),
        Path(r"C:\Windows\Fonts\Kokila.ttf"),
        Path(r"C:\Windows\Fonts\Aparajita.ttf"),
        Path(r"C:\Windows\Fonts\Utsaah.ttf"),
        Path(r"C:\Windows\Fonts\arialuni.ttf"),
        Path("/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf"),
        Path("/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf"),
        Path("/usr/share/fonts/truetype/freefont/FreeSerif.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    for font_dir in [Path(r"C:\Windows\Fonts"), Path("/usr/share/fonts"), Path("/usr/local/share/fonts")]:
        if font_dir.exists():
            patterns = ["*Nirmala*.ttf", "*Mangal*.ttf", "*NotoSansDevanagari*.ttf", "*Kokila*.ttf", "*Aparajita*.ttf", "*Utsaah*.ttf", "*FreeSerif*.ttf"]
            for pattern in patterns:
                candidates.extend(font_dir.rglob(pattern))
    seen = set()
    unique = []
    for item in candidates:
        key = str(item).lower()
        if key not in seen:
            seen.add(key)
            unique.append(item)
    return unique


def _pdf_font_path() -> Path | None:
    for font_path in _font_candidates():
        if font_path.exists():
            return font_path
    return None

def _pdf_unicode_font_name() -> str:
    """Register a system font for Latin/Devanagari PDF output."""
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        font_path = _pdf_font_path()
        if font_path is not None:
            name = "LinguaFusionUnicode"
            if name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(name, str(font_path)))
            return name
    except Exception:
        pass
    return "Helvetica"


def _contains_devanagari(text: str) -> bool:
    return bool(re.search(r"[\u0900-\u097F]", text or ""))


def _wrap_for_image(draw, text: str, font, max_px: int) -> list[str]:
    words = re.split(r"(\s+)", text.strip())
    lines: list[str] = []
    current = ""
    for token in words:
        candidate = current + token
        try:
            width = draw.textbbox((0, 0), candidate, font=font)[2]
        except Exception:
            width = len(candidate) * 12
        if current and width > max_px:
            lines.append(current.strip())
            current = token.strip()
        else:
            current = candidate
    if current.strip():
        lines.append(current.strip())
    out: list[str] = []
    for line in lines:
        try:
            width = draw.textbbox((0, 0), line, font=font)[2]
        except Exception:
            width = len(line) * 12
        if width <= max_px:
            out.append(line)
            continue
        chunk = ""
        for ch in line:
            candidate = chunk + ch
            try:
                cwidth = draw.textbbox((0, 0), candidate, font=font)[2]
            except Exception:
                cwidth = len(candidate) * 12
            if chunk and cwidth > max_px:
                out.append(chunk)
                chunk = ch
            else:
                chunk = candidate
        if chunk:
            out.append(chunk)
    return out or [text]


def _append_pdf_text(story, text: str, style, available_width_pt: float = 470):
    """Append text to a ReportLab story, rendering Devanagari lines as images.

    ReportLab Paragraph can produce black boxes for Devanagari on some Windows
    setups even when a Unicode font is available. Rendering only Devanagari
    lines through Pillow avoids those black boxes while keeping normal text and
    tables as selectable ReportLab text.
    """
    clean = (text or "").strip()
    if not clean:
        return
    if not _contains_devanagari(clean):
        from reportlab.platypus import Paragraph
        from xml.sax.saxutils import escape
        story.append(Paragraph(escape(clean), style))
        return

    font_path = _pdf_font_path()
    if font_path is None:
        from reportlab.platypus import Paragraph
        from xml.sax.saxutils import escape
        story.append(Paragraph(escape(clean), style))
        return

    try:
        from PIL import Image as PILImage, ImageDraw, ImageFont
        from reportlab.platypus import Image as RLImage
        from tempfile import NamedTemporaryFile
        scale = 2
        max_px = int(available_width_pt * scale)
        font_size = max(20, int(float(getattr(style, "fontSize", 10) or 10) * scale * 1.35))
        font = ImageFont.truetype(str(font_path), font_size)
        dummy = PILImage.new("RGB", (max_px, 10), "white")
        draw = ImageDraw.Draw(dummy)
        lines = _wrap_for_image(draw, clean, font, max_px - 20)
        line_h = int(font_size * 1.45)
        img_h = max(line_h + 16, line_h * len(lines) + 16)
        img = PILImage.new("RGB", (max_px, img_h), "white")
        draw = ImageDraw.Draw(img)
        y = 6
        for line in lines:
            draw.text((4, y), line, fill="black", font=font)
            y += line_h
        tmp = NamedTemporaryFile(delete=False, suffix=".png")
        tmp.close()
        img.save(tmp.name)
        story.append(RLImage(tmp.name, width=available_width_pt, height=img_h / scale))
    except Exception:
        from reportlab.platypus import Paragraph
        from xml.sax.saxutils import escape
        story.append(Paragraph(escape(clean), style))


def _resolve_source_language(text: str, source_lang: str) -> str:
    source_lang = normalize_lang(source_lang)
    if source_lang and source_lang != "auto":
        return source_lang

    detected = detect_text_language(text)
    if detected.get("ok") and detected.get("language"):
        return detected["language"]
    return "en"


def _translate_piece(text: str, source_lang: str, target_lang: str, cache: Dict[Tuple[str, str, str], str]) -> str:
    if not text or not text.strip():
        return text

    key = (text, source_lang, target_lang)
    if key in cache:
        return cache[key]

    result = translate_with_views(text, source_lang, target_lang)
    if not result.get("ok"):
        raise RuntimeError(result.get("error") or "Translation failed")

    translated = result.get("translated_text", "")
    cache[key] = translated
    return translated


def _translate_line_preserving_pipes(line: str, source_lang: str, target_lang: str, cache: Dict[Tuple[str, str, str], str]) -> str:
    if not line.strip():
        return line
    if "|" not in line:
        return _translate_piece(line.strip(), source_lang, target_lang, cache)
    if line.strip().replace("|", "").replace("-", "").replace(":", "").strip() == "":
        return line
    leading = line[: len(line) - len(line.lstrip())]
    trailing = line[len(line.rstrip()):]
    cells = line.strip().split("|")
    out_cells = []
    for cell in cells:
        core = cell.strip()
        if not core:
            out_cells.append("")
        elif core in {"---", ":---", "---:", ":---:"}:
            out_cells.append(core)
        else:
            out_cells.append(_translate_piece(core, source_lang, target_lang, cache))
    return leading + " | ".join(out_cells) + trailing


def translate_text_file_preserving_lines(file_path: Path, source_lang: str, target_lang: str, output_suffix: str) -> Path:
    # CSV files are converted to a pipe-table representation first so DOCX/PDF
    # exports can render them as real tables instead of plain paragraphs.
    raw_text = extract_csv_text(file_path) if file_path.suffix.lower() == ".csv" else read_plain_text(file_path)
    resolved_source = _resolve_source_language(raw_text, source_lang)
    cache: Dict[Tuple[str, str, str], str] = {}

    translated_lines = []
    for line in raw_text.splitlines():
        if not line.strip():
            translated_lines.append("")
            continue

        leading = line[: len(line) - len(line.lstrip())]
        trailing = line[len(line.rstrip()):]
        core = line.strip()
        if "|" in core:
            translated_lines.append(_translate_line_preserving_pipes(line, resolved_source, target_lang, cache))
        else:
            translated_lines.append(leading + _translate_piece(core, resolved_source, target_lang, cache) + trailing)

    return _write_translated_text_output("\n".join(translated_lines), output_suffix)


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    """Replace paragraph text while preserving paragraph style and first-run character style."""
    if paragraph.runs:
        first_run = paragraph.runs[0]
        for run in paragraph.runs:
            run.text = ""
        first_run.text = new_text
    else:
        paragraph.add_run(new_text)


def translate_docx_preserving_layout(file_path: Path, source_lang: str, target_lang: str, output_suffix: str) -> Path:
    source_doc = Document(file_path)
    full_text = "\n".join(p.text for p in source_doc.paragraphs if p.text.strip())
    resolved_source = _resolve_source_language(full_text, source_lang)
    cache: Dict[Tuple[str, str, str], str] = {}

    for paragraph in source_doc.paragraphs:
        original = paragraph.text
        if original.strip():
            translated = _translate_piece(original, resolved_source, target_lang, cache)
            _replace_paragraph_text(paragraph, translated)

    for table in source_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original = paragraph.text
                    if original.strip():
                        translated = _translate_piece(original, resolved_source, target_lang, cache)
                        _replace_paragraph_text(paragraph, translated)

    if output_suffix in {".txt", ".pdf"}:
        lines = []
        for paragraph in source_doc.paragraphs:
            if paragraph.text.strip():
                lines.append(paragraph.text)
        for table in source_doc.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    lines.append(" | ".join(cells))
        return _write_translated_text_output("\n".join(lines), output_suffix)

    out = Path(NamedTemporaryFile(delete=False, suffix=".docx").name)
    source_doc.save(out)
    return out


def _is_pipe_table_line(line: str) -> bool:
    stripped = (line or "").strip()
    return "|" in stripped and len([c for c in stripped.split("|") if c.strip()]) >= 2


def _is_pipe_separator(line: str) -> bool:
    stripped = (line or "").strip()
    if not stripped or "|" not in stripped:
        return False
    compact = stripped.replace("|", "").replace("-", "").replace(":", "").replace(" ", "")
    return compact == ""


def _split_pipe_cells(line: str) -> list[str]:
    stripped = (line or "").strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def _blocks_with_pipe_tables(text: str):
    """Yield ('paragraph', str) or ('table', rows) while preserving table runs."""
    lines = (text or "").splitlines()
    i = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            payload = "\n".join(paragraph_buffer).strip()
            paragraph_buffer = []
            if payload:
                return ("paragraph", payload)
        return None

    while i < len(lines):
        line = lines[i]
        if _is_pipe_table_line(line):
            flushed = flush_paragraph()
            if flushed:
                yield flushed
            rows = []
            while i < len(lines) and (_is_pipe_table_line(lines[i]) or _is_pipe_separator(lines[i])):
                if not _is_pipe_separator(lines[i]):
                    rows.append(_split_pipe_cells(lines[i]))
                i += 1
            if rows:
                max_cols = max(len(row) for row in rows)
                rows = [row + [""] * (max_cols - len(row)) for row in rows]
                yield ("table", rows)
            continue
        if line.strip():
            paragraph_buffer.append(line)
        else:
            flushed = flush_paragraph()
            if flushed:
                yield flushed
        i += 1
    flushed = flush_paragraph()
    if flushed:
        yield flushed

def _write_translated_text_output(translated_text: str, output_suffix: str, title: str = "LinguaFusion Translation") -> Path:
    if output_suffix == ".docx":
        out = Path(NamedTemporaryFile(delete=False, suffix=".docx").name)
        doc = Document()
        doc.add_heading(title, level=1)
        for block_type, payload in _blocks_with_pipe_tables(translated_text):
            if block_type == "table":
                rows = payload
                if not rows:
                    continue
                table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx, value in enumerate(row):
                        table.cell(r_idx, c_idx).text = value
                doc.add_paragraph("")
            else:
                for line in str(payload).splitlines():
                    doc.add_paragraph(line)
        doc.save(out)
        return out

    if output_suffix == ".pdf":
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        from reportlab.lib import colors
        from xml.sax.saxutils import escape

        out = Path(NamedTemporaryFile(delete=False, suffix=".pdf").name)
        pdf = SimpleDocTemplate(str(out), pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        unicode_font = _pdf_unicode_font_name()
        for style_name in ["Normal", "BodyText", "Title", "Heading1", "Heading2"]:
            if style_name in styles:
                styles[style_name].fontName = unicode_font
        story = [Paragraph(escape(title), styles["Title"]), Spacer(1, 0.35*cm)]
        for block_type, payload in _blocks_with_pipe_tables(translated_text):
            if block_type == "table":
                rows = payload
                table_data = [[Paragraph(escape(str(cell)), styles["BodyText"]) for cell in row] for row in rows]
                table = Table(table_data, repeatRows=1)
                table.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.grey),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ]))
                story.append(table)
                story.append(Spacer(1, 0.25*cm))
            else:
                for line in str(payload).splitlines():
                    if line.strip():
                        _append_pdf_text(story, line.strip(), styles["BodyText"])
                        story.append(Spacer(1, 0.12*cm))
                story.append(Spacer(1, 0.12*cm))
        pdf.build(story)
        return out

    out = Path(NamedTemporaryFile(delete=False, suffix=".txt").name)
    out.write_text(translated_text, encoding="utf-8")
    return out

def translate_extracted_text_preserving_blocks(text: str, source_lang: str, target_lang: str, output_suffix: str, title: str = "LinguaFusion Translation") -> Path:
    resolved_source = _resolve_source_language(text, source_lang)
    cache: Dict[Tuple[str, str, str], str] = {}
    translated_blocks = []
    for block in text.split("\n\n"):
        if not block.strip():
            translated_blocks.append("")
            continue
        translated_lines = []
        for line in block.splitlines():
            if not line.strip():
                translated_lines.append("")
            else:
                leading = line[: len(line) - len(line.lstrip())]
                trailing = line[len(line.rstrip()):]
                if "|" in line:
                    translated_lines.append(_translate_line_preserving_pipes(line, resolved_source, target_lang, cache))
                else:
                    translated_lines.append(leading + _translate_piece(line.strip(), resolved_source, target_lang, cache) + trailing)
        translated_blocks.append("\n".join(translated_lines))
    return _write_translated_text_output("\n\n".join(translated_blocks), output_suffix, title=title)


def translate_document_preserving_format(file_path: Path, source_lang: str, target_lang: str, output_format: str) -> Path:
    output_format = output_format.lower().strip().lstrip(".")
    if output_format not in {"txt", "docx", "pdf"}:
        raise ValueError("Format-preserving export currently supports TXT, DOCX and PDF.")

    suffix = file_path.suffix.lower()
    output_suffix = f".{output_format}"

    if suffix == ".docx":
        return translate_docx_preserving_layout(file_path, source_lang, target_lang, output_suffix)

    if suffix in {".txt", ".md", ".csv"}:
        return translate_text_file_preserving_lines(file_path, source_lang, target_lang, output_suffix)

    if suffix == ".pdf":
        extracted = extract_text_from_file(file_path, source_lang)
        if not extracted.get("ok"):
            raise ValueError(extracted.get("error") or "PDF text extraction failed.")
        return translate_extracted_text_preserving_blocks(extracted.get("text", ""), source_lang, target_lang, output_suffix, title="LinguaFusion PDF Translation")

    raise ValueError(f"Format-preserving export is not available for {suffix} yet. Use normal export for this file type.")
