from __future__ import annotations

import csv
import json
import tempfile
from pathlib import Path
from typing import Any

try:
    from docx import Document
except Exception as exc:  # pragma: no cover - depends on local install
    Document = None
    DOCX_IMPORT_ERROR = exc
else:
    DOCX_IMPORT_ERROR = None

try:
    import fitz  # PyMuPDF
except Exception as exc:  # pragma: no cover
    fitz = None
    FITZ_IMPORT_ERROR = exc
else:
    FITZ_IMPORT_ERROR = None

try:
    from bs4 import BeautifulSoup
except Exception as exc:  # pragma: no cover
    BeautifulSoup = None
    BS4_IMPORT_ERROR = exc
else:
    BS4_IMPORT_ERROR = None

try:
    from striprtf.striprtf import rtf_to_text
except Exception as exc:  # pragma: no cover
    rtf_to_text = None
    RTF_IMPORT_ERROR = exc
else:
    RTF_IMPORT_ERROR = None

try:
    from PIL import Image
except Exception as exc:  # pragma: no cover
    Image = None
    PIL_IMPORT_ERROR = exc
else:
    PIL_IMPORT_ERROR = None

try:
    import pytesseract
except Exception as exc:  # pragma: no cover
    pytesseract = None
    TESSERACT_IMPORT_ERROR = exc
else:
    TESSERACT_IMPORT_ERROR = None

from backend.config.paths import TESSERACT_EXE

if pytesseract is not None and TESSERACT_EXE.exists():
    pytesseract.pytesseract.tesseract_cmd = str(TESSERACT_EXE)

SUPPORTED_TEXT_EXTENSIONS = {
    ".txt", ".md", ".csv", ".json", ".xml",
    ".html", ".htm", ".rtf", ".docx", ".pdf",
}

OCR_LANGS = {
    "auto": "eng+deu+spa+hin",
    "en": "eng",
    "de": "deu",
    "es": "spa",
    "hi": "hin",
}


def _missing_dependency(name: str, exc: Exception | None) -> RuntimeError:
    detail = f": {exc}" if exc else ""
    return RuntimeError(f"Missing optional dependency for this file type: {name}{detail}")


def read_plain_text(file_path: Path) -> str:
    try:
        return file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return file_path.read_text(encoding="latin-1")


def extract_docx_text(file_path: Path) -> str:
    if Document is None:
        raise _missing_dependency("python-docx", DOCX_IMPORT_ERROR)

    document = Document(file_path)
    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (getattr(paragraph.style, "name", "") or "").lower()
        if style_name.startswith("heading"):
            parts.append(f"\n## {text}\n")
        else:
            parts.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        table_lines = [f"\nTable {table_index}"]
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                cells.append(cell_text)
            if any(cell.strip() for cell in cells):
                table_lines.append(" | ".join(cells))
        if len(table_lines) > 1:
            parts.append("\n".join(table_lines))

    return "\n\n".join(part.strip() for part in parts if part.strip())


def extract_pdf_embedded_text(file_path: Path) -> str:
    if fitz is None:
        raise _missing_dependency("pymupdf", FITZ_IMPORT_ERROR)

    parts = []
    with fitz.open(file_path) as doc:
        for page_number, page in enumerate(doc, start=1):
            blocks = page.get_text("blocks") or []
            page_parts = []
            for block in sorted(blocks, key=lambda item: (round(item[1], 1), round(item[0], 1))):
                if len(block) >= 5:
                    text = str(block[4]).strip()
                    if text:
                        page_parts.append(text)
            if not page_parts:
                text = page.get_text("text").strip()
                if text:
                    page_parts.append(text)
            if page_parts:
                parts.append(f"\n--- Page {page_number} ---\n" + "\n".join(page_parts))
    return "\n\n".join(parts).strip()


def extract_pdf_ocr_text(file_path: Path, lang: str = "en") -> str:
    if fitz is None:
        raise _missing_dependency("pymupdf", FITZ_IMPORT_ERROR)
    if Image is None:
        raise _missing_dependency("pillow", PIL_IMPORT_ERROR)
    if pytesseract is None:
        raise _missing_dependency("pytesseract", TESSERACT_IMPORT_ERROR)

    tesseract_lang = OCR_LANGS.get(lang, "eng")
    parts = []
    with fitz.open(file_path) as doc:
        for page_number, page in enumerate(doc, start=1):
            pix = page.get_pixmap(dpi=250)
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_img:
                temp_img_path = Path(temp_img.name)
            pix.save(temp_img_path)
            try:
                with Image.open(temp_img_path) as image:
                    text = pytesseract.image_to_string(image, lang=tesseract_lang).strip()
            finally:
                try:
                    temp_img_path.unlink()
                except Exception:
                    pass
            if text:
                parts.append(f"\n--- OCR Page {page_number} ---\n{text}")
    return "\n".join(parts).strip()


def extract_pdf_text(file_path: Path, lang: str = "en") -> dict:
    embedded_text = extract_pdf_embedded_text(file_path)
    if embedded_text.strip():
        return {"ok": True, "text": embedded_text, "file_type": ".pdf", "method": "embedded_text", "error": None}

    ocr_text = extract_pdf_ocr_text(file_path, lang)
    if ocr_text.strip():
        return {"ok": True, "text": ocr_text, "file_type": ".pdf", "method": "ocr_fallback", "error": None}

    return {"ok": False, "text": "", "file_type": ".pdf", "method": "none", "error": "No readable text found, including OCR fallback."}


def extract_html_text(file_path: Path) -> str:
    if BeautifulSoup is None:
        raise _missing_dependency("beautifulsoup4", BS4_IMPORT_ERROR)
    html = read_plain_text(file_path)
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)


def extract_rtf_text(file_path: Path) -> str:
    if rtf_to_text is None:
        raise _missing_dependency("striprtf", RTF_IMPORT_ERROR)
    return rtf_to_text(read_plain_text(file_path))


def extract_json_text(file_path: Path) -> str:
    raw = read_plain_text(file_path)
    try:
        data: Any = json.loads(raw)
        return json.dumps(data, indent=2, ensure_ascii=False)
    except Exception:
        return raw


def extract_csv_text(file_path: Path) -> str:
    rows = []
    with open(file_path, newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append([cell.strip() for cell in row])
    if not rows:
        return ""
    max_cols = max(len(row) for row in rows)
    normalized = [row + [""] * (max_cols - len(row)) for row in rows]
    lines = [" | ".join(row) for row in normalized]
    if len(normalized) >= 2 and max_cols > 1:
        separator = " | ".join(["---"] * max_cols)
        lines.insert(1, separator)
    return "\n".join(lines)


def extract_text_from_file(file_path: Path, lang: str = "en") -> dict:
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_TEXT_EXTENSIONS:
        return {"ok": False, "text": "", "file_type": suffix, "method": "unsupported", "error": f"Unsupported reader file type: {suffix}"}

    try:
        if suffix == ".docx":
            text, method = extract_docx_text(file_path), "docx"
        elif suffix == ".pdf":
            return extract_pdf_text(file_path, lang)
        elif suffix in {".html", ".htm"}:
            text, method = extract_html_text(file_path), "html"
        elif suffix == ".rtf":
            text, method = extract_rtf_text(file_path), "rtf"
        elif suffix == ".json":
            text, method = extract_json_text(file_path), "json"
        elif suffix == ".csv":
            text, method = extract_csv_text(file_path), "csv"
        else:
            text, method = read_plain_text(file_path), "plain_text"

        if not text.strip():
            return {"ok": False, "text": "", "file_type": suffix, "method": method, "error": "No readable text found."}
        return {"ok": True, "text": text, "file_type": suffix, "method": method, "error": None}
    except Exception as exc:
        return {"ok": False, "text": "", "file_type": suffix, "method": "error", "error": str(exc)}
