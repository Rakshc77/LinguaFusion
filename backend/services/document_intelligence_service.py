from __future__ import annotations

import re
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Any, Dict, List
from xml.sax.saxutils import escape

from docx import Document

WORD_RE = re.compile(r"[\w\u0900-\u097F]+(?:[-'][\w\u0900-\u097F]+)?", re.UNICODE)
SENTENCE_END_RE = re.compile(r"(?<=[.!?।])\s+(?=[A-ZÄÖÜ¿¡\u0900-\u097F0-9])")
TECH_TERM_RE = re.compile(
    r"\b(?:[A-Z]{2,}(?:-[A-Z0-9]+)*|LiDAR|MATLAB|Python|Wireless\s+InSite|PDP|RMS\s+Delay\s+Spread|DOCX|PDF|OCR|TTS|STT|VAD|LFIE|Piper|Whisper)\b"
)




def _font_candidates() -> list[Path]:
    candidates = [
        Path(r"C:\Windows\Fonts\Nirmala.ttf"),
        Path(r"C:\Windows\Fonts\NirmalaB.ttf"),
        Path(r"C:\Windows\Fonts\Mangal.ttf"),
        Path(r"C:\Windows\Fonts\Kokila.ttf"),
        Path(r"C:\Windows\Fonts\Aparajita.ttf"),
        Path(r"C:\Windows\Fonts\Utsaah.ttf"),
        Path(r"C:\Windows\Fonts\arialuni.ttf"),
        Path("/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf"),
        Path("/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf"),
        Path("/usr/share/fonts/truetype/freefont/FreeSerif.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    for font_dir in [Path(r"C:\Windows\Fonts"), Path("/usr/share/fonts"), Path("/usr/local/share/fonts")]:
        if font_dir.exists():
            patterns = ["*Nirmala*.ttf", "*Mangal*.ttf", "*NotoSansDevanagari*.ttf", "*Kokila*.ttf", "*Aparajita*.ttf", "*Utsaah*.ttf", "*FreeSerif*.ttf"]
            for pattern in patterns:
                candidates.extend(font_dir.rglob(pattern))
    seen = set()
    unique = []
    for item in candidates:
        key = str(item).lower()
        if key not in seen:
            seen.add(key)
            unique.append(item)
    return unique


def _pdf_font_path() -> Path | None:
    for font_path in _font_candidates():
        if font_path.exists():
            return font_path
    return None

def _pdf_unicode_font_name() -> str:
    """Register a system font for Latin/Devanagari PDF output."""
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        font_path = _pdf_font_path()
        if font_path is not None:
            name = "LinguaFusionUnicode"
            if name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(name, str(font_path)))
            return name
    except Exception:
        pass
    return "Helvetica"


def _contains_devanagari(text: str) -> bool:
    return bool(re.search(r"[\u0900-\u097F]", text or ""))


def _wrap_for_image(draw, text: str, font, max_px: int) -> list[str]:
    words = re.split(r"(\s+)", text.strip())
    lines: list[str] = []
    current = ""
    for token in words:
        candidate = current + token
        try:
            width = draw.textbbox((0, 0), candidate, font=font)[2]
        except Exception:
            width = len(candidate) * 12
        if current and width > max_px:
            lines.append(current.strip())
            current = token.strip()
        else:
            current = candidate
    if current.strip():
        lines.append(current.strip())
    out: list[str] = []
    for line in lines:
        try:
            width = draw.textbbox((0, 0), line, font=font)[2]
        except Exception:
            width = len(line) * 12
        if width <= max_px:
            out.append(line)
            continue
        chunk = ""
        for ch in line:
            candidate = chunk + ch
            try:
                cwidth = draw.textbbox((0, 0), candidate, font=font)[2]
            except Exception:
                cwidth = len(candidate) * 12
            if chunk and cwidth > max_px:
                out.append(chunk)
                chunk = ch
            else:
                chunk = candidate
        if chunk:
            out.append(chunk)
    return out or [text]


def _append_pdf_text(story, text: str, style, available_width_pt: float = 470):
    """Append text to a ReportLab story, rendering Devanagari lines as images.

    ReportLab Paragraph can produce black boxes for Devanagari on some Windows
    setups even when a Unicode font is available. Rendering only Devanagari
    lines through Pillow avoids those black boxes while keeping normal text and
    tables as selectable ReportLab text.
    """
    clean = (text or "").strip()
    if not clean:
        return
    if not _contains_devanagari(clean):
        from reportlab.platypus import Paragraph
        from xml.sax.saxutils import escape
        story.append(Paragraph(escape(clean), style))
        return

    font_path = _pdf_font_path()
    if font_path is None:
        from reportlab.platypus import Paragraph
        from xml.sax.saxutils import escape
        story.append(Paragraph(escape(clean), style))
        return

    try:
        from PIL import Image as PILImage, ImageDraw, ImageFont
        from reportlab.platypus import Image as RLImage
        from tempfile import NamedTemporaryFile
        scale = 2
        max_px = int(available_width_pt * scale)
        font_size = max(20, int(float(getattr(style, "fontSize", 10) or 10) * scale * 1.35))
        font = ImageFont.truetype(str(font_path), font_size)
        dummy = PILImage.new("RGB", (max_px, 10), "white")
        draw = ImageDraw.Draw(dummy)
        lines = _wrap_for_image(draw, clean, font, max_px - 20)
        line_h = int(font_size * 1.45)
        img_h = max(line_h + 16, line_h * len(lines) + 16)
        img = PILImage.new("RGB", (max_px, img_h), "white")
        draw = ImageDraw.Draw(img)
        y = 6
        for line in lines:
            draw.text((4, y), line, fill="black", font=font)
            y += line_h
        tmp = NamedTemporaryFile(delete=False, suffix=".png")
        tmp.close()
        img.save(tmp.name)
        story.append(RLImage(tmp.name, width=available_width_pt, height=img_h / scale))
    except Exception:
        from reportlab.platypus import Paragraph
        from xml.sax.saxutils import escape
        story.append(Paragraph(escape(clean), style))


def format_duration_label(seconds: int | float) -> str:
    """Return a commercial-looking duration such as '2 min 41 sec'."""
    try:
        total = max(0, int(round(float(seconds))))
    except Exception:
        total = 0
    minutes, secs = divmod(total, 60)
    if minutes and secs:
        return f"{minutes} min {secs} sec"
    if minutes:
        return f"{minutes} min"
    return f"{secs} sec"


def normalize_document_text(text: str) -> str:
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\t\f\v]+", " ", text)
    text = re.sub(r"[ \u00a0]{2,}", " ", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def split_paragraphs(text: str) -> List[str]:
    clean = normalize_document_text(text)
    if not clean:
        return []
    return [p.strip() for p in re.split(r"\n\s*\n", clean) if p.strip()]


def split_sentences_with_ranges(text: str) -> List[Dict[str, Any]]:
    """Return sentence/line ranges against the original text for UI highlighting/seeking.

    Phase 3 beta3: document text often starts with headings, numbered sections,
    table labels and short instruction lines without terminal punctuation. The
    earlier punctuation-only splitter skipped those blocks, so playback highlight
    was offset until the first normal paragraph. This splitter treats every
    non-empty line as a boundary first, then splits long prose lines by sentence
    punctuation.
    """
    original = text or ""
    ranges: List[Dict[str, Any]] = []
    length = len(original)

    def add_range(start: int, end: int, sentence_text: str) -> None:
        sentence = (sentence_text or "").strip()
        if not sentence:
            return
        ranges.append({
            "index": len(ranges),
            "start": max(0, min(start, length)),
            "end": max(0, min(end, length)),
            "text": sentence,
            "words": len(WORD_RE.findall(sentence)),
        })

    for line_match in re.finditer(r"[^\n]+", original, flags=re.UNICODE):
        line = line_match.group(0)
        stripped = line.strip()
        if not stripped:
            continue
        line_start = line_match.start() + (len(line) - len(line.lstrip()))
        line_end = line_match.end() - (len(line) - len(line.rstrip()))

        # Very short headings/table rows are better as one highlight unit.
        if len(stripped) <= 90 and not re.search(r"[.!?।]", stripped):
            add_range(line_start, line_end, stripped)
            continue

        found_piece = False
        for match in re.finditer(r"[^.!?।]+(?:[.!?।]+|$)", stripped, flags=re.UNICODE):
            raw = match.group(0)
            piece = raw.strip()
            if not piece:
                continue
            found_piece = True
            piece_start = line_start + match.start() + (len(raw) - len(raw.lstrip()))
            piece_end = line_start + match.end() - (len(raw) - len(raw.rstrip()))
            add_range(piece_start, piece_end, piece)

        if not found_piece:
            add_range(line_start, line_end, stripped)

    if not ranges and original.strip():
        stripped_start = len(original) - len(original.lstrip())
        stripped_end = len(original.rstrip())
        add_range(stripped_start, stripped_end, original.strip())

    return ranges


def detect_headings(text: str) -> List[Dict[str, Any]]:
    headings: List[Dict[str, Any]] = []
    position = 0
    for line in (text or "").splitlines():
        stripped = line.strip()
        if not stripped:
            position += len(line) + 1
            continue
        is_heading = False
        level = 2
        if stripped.startswith("#"):
            is_heading = True
            level = min(6, max(1, len(stripped) - len(stripped.lstrip("#"))))
            title = stripped.lstrip("#").strip()
        elif re.match(r"^(?:\d+[.)]|[A-Z]\.)\s+", stripped) and len(stripped) <= 90:
            is_heading = True
            title = stripped
        elif len(stripped) <= 75 and not re.search(r"[.!?।]$", stripped) and (stripped.istitle() or stripped.isupper()):
            is_heading = True
            title = stripped
        else:
            title = stripped
        if is_heading:
            headings.append({"title": title, "level": level, "position": position})
        position += len(line) + 1
    return headings[:50]


def analyze_document(text: str) -> Dict[str, Any]:
    normalized = normalize_document_text(text)
    words = WORD_RE.findall(normalized)
    sentences = split_sentences_with_ranges(normalized)
    paragraphs = split_paragraphs(normalized)
    technical_terms = []
    seen = set()
    for match in TECH_TERM_RE.finditer(normalized):
        term = re.sub(r"\s+", " ", match.group(0)).strip()
        key = term.lower()
        if key not in seen:
            technical_terms.append(term)
            seen.add(key)
    reading_wpm = 165
    speaking_wpm = 145
    estimated_reading_seconds = int(round((len(words) / reading_wpm) * 60)) if words else 0
    estimated_speaking_seconds = int(round((len(words) / speaking_wpm) * 60)) if words else 0
    return {
        "ok": True,
        "characters": len(normalized),
        "words": len(words),
        "sentences": len(sentences),
        "paragraphs": len(paragraphs),
        "line_count": len(normalized.splitlines()) if normalized else 0,
        "estimated_reading_minutes": round(estimated_reading_seconds / 60, 1) if words else 0.0,
        "estimated_speaking_minutes": round(estimated_speaking_seconds / 60, 1) if words else 0.0,
        "estimated_reading_seconds": estimated_reading_seconds,
        "estimated_speaking_seconds": estimated_speaking_seconds,
        "estimated_speaking_label": format_duration_label(estimated_speaking_seconds),
        "technical_terms": technical_terms[:40],
        "headings": detect_headings(normalized),
        "sentence_ranges": sentences,
    }


def export_reader_document(text: str, output_format: str, title: str = "LinguaFusion Reader Export") -> Path:
    output_format = (output_format or "txt").lower().strip().lstrip(".")
    text = normalize_document_text(text)
    if output_format == "txt":
        out = Path(NamedTemporaryFile(delete=False, suffix=".txt").name)
        out.write_text(text + "\n", encoding="utf-8")
        return out
    if output_format == "docx":
        out = Path(NamedTemporaryFile(delete=False, suffix=".docx").name)
        doc = Document()
        doc.add_heading(title, level=1)
        for paragraph in split_paragraphs(text) or [text]:
            doc.add_paragraph(paragraph)
        doc.save(out)
        return out
    if output_format == "pdf":
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        out = Path(NamedTemporaryFile(delete=False, suffix=".pdf").name)
        doc = SimpleDocTemplate(str(out), pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        unicode_font = _pdf_unicode_font_name()
        for style_name in ["Normal", "BodyText", "Title", "Heading1", "Heading2"]:
            if style_name in styles:
                styles[style_name].fontName = unicode_font
        story = [Paragraph(escape(title), styles["Title"]), Spacer(1, 0.35*cm)]
        for paragraph in split_paragraphs(text) or [text]:
            for line in paragraph.splitlines():
                if line.strip():
                    _append_pdf_text(story, line.strip(), styles["BodyText"])
                    story.append(Spacer(1, 0.15*cm))
            story.append(Spacer(1, 0.15*cm))
        doc.build(story)
        return out
    raise ValueError("Reader export supports TXT, DOCX and PDF only.")
