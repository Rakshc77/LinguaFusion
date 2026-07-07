from __future__ import annotations

from pathlib import Path
from tempfile import NamedTemporaryFile
import sys

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document

from backend.services.document_intelligence_service import analyze_document, export_reader_document
from backend.services.file_reader_service import extract_text_from_file, extract_csv_text


def test_document_analysis():
    text = """LinguaFusion Document Test\n\nThis reader paragraph mentions MATLAB, Python and Wireless InSite.\nHallo zusammen, wir testen Zeilenumbrüche und technische Begriffe."""
    analysis = analyze_document(text)
    assert analysis["ok"] is True
    assert analysis["words"] >= 12
    assert analysis["sentences"] >= 2
    assert "MATLAB" in analysis["technical_terms"]
    assert analysis["sentence_ranges"]


def test_docx_table_extraction():
    path = Path(NamedTemporaryFile(delete=False, suffix=".docx").name)
    doc = Document()
    doc.add_heading("Reader Export Test", level=1)
    doc.add_paragraph("Rajarshi lives in Melsungen.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Tool"
    table.cell(0, 1).text = "Expected"
    table.cell(1, 0).text = "Wireless InSite"
    table.cell(1, 1).text = "Preserve name"
    doc.save(path)
    result = extract_text_from_file(path, "en")
    assert result["ok"] is True
    assert "## Reader Export Test" in result["text"]
    assert "Table 1" in result["text"]
    assert "Wireless InSite | Preserve name" in result["text"]
    path.unlink(missing_ok=True)


def test_reader_exports():
    text = "Line one.\n\nLine two with MATLAB and Python."
    for fmt in ["txt", "docx", "pdf"]:
        out = export_reader_document(text, fmt, title="Phase 3 Test")
        assert out.exists()
        assert out.stat().st_size > 0
        out.unlink(missing_ok=True)


def test_csv_layout():
    path = Path(NamedTemporaryFile(delete=False, suffix=".csv").name)
    path.write_text("Item,Value\nCity,Melsungen\nTool,Wireless InSite\n", encoding="utf-8")
    text = extract_csv_text(path)
    assert "Item | Value" in text
    assert "--- | ---" in text
    assert "Tool | Wireless InSite" in text
    path.unlink(missing_ok=True)


if __name__ == "__main__":
    test_document_analysis()
    test_docx_table_extraction()
    test_reader_exports()
    test_csv_layout()
    print("Phase 3 document workflow tests passed")
