import os
import re
import subprocess
import sys
import json
import hashlib
import math
import time
import requests
import pygame
import numpy as np
import sounddevice as sd
import soundfile as sf
from pydub import AudioSegment
from pathlib import Path
from uuid import uuid4
from concurrent.futures import ThreadPoolExecutor

from PySide6.QtCore import Qt, QTimer, QPointF, Signal
from PySide6.QtGui import QShortcut, QKeySequence, QPainter, QPen, QColor, QPainterPath, QTextCursor, QTextCharFormat
from PySide6.QtWidgets import (
    QApplication, QComboBox, QFileDialog, QFrame, QGridLayout, QHBoxLayout,
    QLabel, QLineEdit, QMainWindow, QPushButton, QSizePolicy, QStackedWidget, QTextEdit,
    QVBoxLayout, QWidget, QCheckBox, QInputDialog, QMessageBox
)

SERVER_URL = "http://localhost:8000"

LANGUAGES = [
    ("English", "en"),
    ("German", "de"),
    ("Spanish", "es"),
    ("Hindi", "hi"),
]

DOCUMENT_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".rtf", ".html", ".htm", ".csv", ".json", ".xml"}
AUDIO_EXTENSIONS = {".wav", ".mp3", ".m4a", ".ogg", ".flac", ".aac", ".wma"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}
OCR_EXTENSIONS = IMAGE_EXTENSIONS | {".pdf"}


class Card(QFrame):
    def __init__(self, object_name="Card"):
        super().__init__()
        self.setObjectName(object_name)
        self.setFrameShape(QFrame.NoFrame)


class AudioWaveform(QWidget):
    """Spotify/Apple-style waveform drawn from the actual WAV audio."""
    seek_requested = Signal(float)

    def __init__(self):
        super().__init__()
        self.progress = 0.0
        self.peaks = []
        self.setMinimumHeight(62)

    def set_progress(self, progress: float):
        self.progress = max(0.0, min(1.0, progress))
        self.update()

    def clear_waveform(self):
        self.peaks = []
        self.progress = 0.0
        self.update()

    def _emit_seek_from_x(self, x_position: int):
        width = max(self.width(), 1)
        progress = max(0.0, min(1.0, x_position / width))
        self.seek_requested.emit(progress)

    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self._emit_seek_from_x(event.position().x())

    def mouseMoveEvent(self, event):
        if event.buttons() & Qt.LeftButton:
            self._emit_seek_from_x(event.position().x())

    def load_audio(self, audio_path: str, bins: int = 420):
        try:
            audio = AudioSegment.from_file(audio_path)
            audio = audio.set_channels(1)
            samples = audio.get_array_of_samples()

            if not samples:
                self.peaks = []
                self.update()
                return

            total = len(samples)
            bins = max(80, min(bins, 900))
            step = max(1, total // bins)
            raw = []

            for i in range(0, total, step):
                chunk = samples[i:i + step]
                if not chunk:
                    continue
                peak = max(abs(x) for x in chunk)
                raw.append(float(peak))

            maximum = max(raw) if raw else 1.0
            normalized = [value / maximum for value in raw]

            # Smooth the contour so it looks like a podcast waveform, not blocks.
            smoothed = []
            window = 5
            for i in range(len(normalized)):
                left = max(0, i - window)
                right = min(len(normalized), i + window + 1)
                smoothed.append(sum(normalized[left:right]) / (right - left))

            self.peaks = smoothed
            self.progress = 0.0
            self.update()
        except Exception:
            self.peaks = []
            self.update()

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing, True)

        width = max(self.width(), 1)
        height = max(self.height(), 1)
        center_y = height / 2
        max_amp = height * 0.36

        baseline_pen = QPen(QColor(54, 72, 104), 2)
        inactive_pen = QPen(QColor(70, 88, 120), 2)
        active_pen = QPen(QColor(96, 120, 255), 3)
        playhead_pen = QPen(QColor(245, 248, 255), 2)

        painter.setPen(baseline_pen)
        painter.drawLine(0, int(center_y), width, int(center_y))

        if not self.peaks:
            # Calm empty-state curve.
            path = QPainterPath()
            path.moveTo(0, center_y)
            for x in range(width + 1):
                t = x / width
                y = center_y + math.sin(t * math.pi * 10) * max_amp * 0.25
                path.lineTo(x, y)
            painter.setPen(inactive_pen)
            painter.drawPath(path)
            return

        top_path = QPainterPath()
        bottom_path = QPainterPath()

        count = len(self.peaks)
        for i, peak in enumerate(self.peaks):
            x = (i / max(count - 1, 1)) * width
            amp = max(2.0, peak * max_amp)
            y_top = center_y - amp
            y_bottom = center_y + amp
            if i == 0:
                top_path.moveTo(x, y_top)
                bottom_path.moveTo(x, y_bottom)
            else:
                top_path.lineTo(x, y_top)
                bottom_path.lineTo(x, y_bottom)

        painter.setPen(inactive_pen)
        painter.drawPath(top_path)
        painter.drawPath(bottom_path)

        active_width = int(width * self.progress)
        if active_width > 0:
            painter.save()
            painter.setClipRect(0, 0, active_width, height)
            painter.setPen(active_pen)
            painter.drawPath(top_path)
            painter.drawPath(bottom_path)
            painter.restore()

        # Smooth podcast-style progress line.
        progress_y = int(height - 7)
        painter.setPen(QPen(QColor(54, 72, 104), 2))
        painter.drawLine(0, progress_y, width, progress_y)
        painter.setPen(QPen(QColor(59, 130, 246), 4))
        painter.drawLine(0, progress_y, active_width, progress_y)

        x = active_width
        painter.setPen(playhead_pen)
        painter.drawLine(x, 5, x, height - 2)
        painter.setBrush(QColor(245, 248, 255))
        painter.drawEllipse(QPointF(x, 8), 5, 5)


class LinguaFusionWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("LinguaFusion")
        self.resize(1500, 940)
        self.setAcceptDrops(True)

        self.executor = ThreadPoolExecutor(max_workers=4)
        self.generated_audio_files = []
        self.audio_paused = False
        self.current_page_name = "Translate"
        self.recent_history = {name: [] for name in ["Translate", "Reader", "Speech", "OCR", "Notes", "Settings"]}
        self.reader_current_file = "None"
        self.reader_detected_language = "Auto"
        self.reader_voice_label = "Auto"
        self.reader_playback_status = "Stopped"
        self.reader_base_audio_path = None
        self.reader_current_audio_path = None
        self.reader_active_audio_path = None
        self.reader_audio_duration_ms = 0
        self.reader_active_duration_ms = 0
        self.reader_last_rate = 1.0
        self.reader_paused_pos_ms = 0
        self.reader_seek_base_original_ms = 0
        self.reader_waveform_bars = ""
        self.reader_analysis = {}
        self.reader_sentence_ranges = []
        self.reader_current_sentence_index = -1
        self.reader_bookmarks = []
        self.reader_bookmark_counter = 0
        self.reader_cursor_audio_cache = {}
        self.reader_full_audio_cache_key = None
        self.reader_full_audio_path = None
        self.reader_highlight_enabled = True
        self.reader_playback_sentence_offset = 0
        self.reader_playback_sentence_count = 0
        self.translate_document_name = None
        self.translate_document_source_path = None
        self.last_translation_plain_text = ""
        self.last_translation_tts_text = ""
        self.last_translation_input_snapshot = ""
        self.last_translation_source_snapshot = ""
        self.last_translation_target_snapshot = ""
        self.active_audio_context = "Reader"

        self.speech_is_recording = False
        self.speech_is_paused = False
        self.speech_frames = []
        self.speech_levels = []
        self.speech_stream = None
        self.speech_sample_rate = 16000
        self.speech_channels = 1
        self.speech_recording_started_at = None
        self.speech_elapsed_when_paused = 0.0
        self.speech_current_wav_path = None
        self.last_speech_transcript = ""
        self.last_speech_translation = ""
        self.last_speech_language = "auto"
        self.speech_live_transcribing = False
        self.speech_last_live_frame_count = 0
        self.speech_auto_final_pending = False
        self.corrections = {}
        self.smart_provider_status = {}

        self.audio_backend_ready = False
        self.audio_backend_error = None
        if os.name == "nt":
            os.environ.setdefault("SDL_AUDIODRIVER", "directsound")
        try:
            pygame.mixer.init()
            self.audio_backend_ready = True
        except Exception as exc:
            self.audio_backend_error = str(exc)

        self.playback_timer = QTimer(self)
        self.playback_timer.setInterval(16)
        self.playback_timer.timeout.connect(self.update_playback_ui)

        self.speech_timer = QTimer(self)
        self.speech_timer.setInterval(80)
        self.speech_timer.timeout.connect(self.update_speech_recording_ui)

        self.speech_live_timer = QTimer(self)
        self.speech_live_timer.setInterval(6000)
        self.speech_live_timer.timeout.connect(self.live_transcribe_speech_snapshot)

        self.build_shell()
        self.load_corrections_from_backend()
        self.load_smart_provider_status()
        self.apply_style()
        self.switch_page("Translate")
        self.check_health()

    # ---------- Shell ----------
    def build_shell(self):
        root = QWidget()
        root.setObjectName("Root")
        root_layout = QVBoxLayout(root)
        root_layout.setContentsMargins(0, 0, 0, 0)
        root_layout.setSpacing(0)

        content = QWidget()
        content_layout = QHBoxLayout(content)
        content_layout.setContentsMargins(0, 0, 0, 0)
        content_layout.setSpacing(0)

        self.sidebar = self.build_sidebar()
        self.main_area = self.build_main_area()
        self.right_panel = self.build_right_panel()

        content_layout.addWidget(self.sidebar)
        content_layout.addWidget(self.main_area, 1)
        content_layout.addWidget(self.right_panel)

        self.footer = self.build_footer()

        root_layout.addWidget(content, 1)
        root_layout.addWidget(self.footer)
        self.setCentralWidget(root)

    def build_sidebar(self):
        sidebar = QFrame()
        sidebar.setObjectName("Sidebar")
        sidebar.setFixedWidth(288)
        layout = QVBoxLayout(sidebar)
        layout.setContentsMargins(26, 22, 24, 22)
        layout.setSpacing(16)

        brand = QHBoxLayout()
        logo = QLabel("💬")
        logo.setObjectName("Logo")
        title_col = QVBoxLayout()
        title = QLabel("LinguaFusion")
        title.setObjectName("AppTitle")
        subtitle = QLabel("All-in-one Language Assistant")
        subtitle.setObjectName("Muted")
        title_col.addWidget(title)
        title_col.addWidget(subtitle)
        brand.addWidget(logo)
        brand.addLayout(title_col)
        layout.addLayout(brand)
        layout.addSpacing(22)

        self.nav_buttons = {}
        nav_items = [
            ("Translate", "🌐"),
            ("Reader", "📖"),
            ("Speech", "🎙"),
            ("OCR", "▣"),
            ("Notes", "📝"),
            ("Settings", "⚙"),
        ]
        for name, icon in nav_items:
            btn = QPushButton(f"{icon}   {name}")
            btn.setObjectName("NavButton")
            btn.setCheckable(True)
            btn.clicked.connect(lambda checked=False, page=name: self.switch_page(page))
            self.nav_buttons[name] = btn
            layout.addWidget(btn)

        layout.addStretch(1)

        recent_header = QHBoxLayout()
        recent = QLabel("RECENT")
        recent.setObjectName("SectionLabel")
        more = QLabel("More")
        more.setObjectName("LinkLabel")
        recent_header.addWidget(recent)
        recent_header.addStretch(1)
        recent_header.addWidget(more)
        layout.addLayout(recent_header)

        self.recent_items_box = QVBoxLayout()
        layout.addLayout(self.recent_items_box)
        self.update_recent_items("Translate")

        import_btn = QPushButton("＋  Import File")
        import_btn.setObjectName("PrimaryButton")
        import_btn.clicked.connect(lambda: self.switch_page("Reader"))
        layout.addWidget(import_btn)
        return sidebar

    def build_main_area(self):
        area = QWidget()
        area.setObjectName("MainArea")
        layout = QVBoxLayout(area)
        layout.setContentsMargins(28, 20, 12, 14)
        layout.setSpacing(18)

        topbar = QHBoxLayout()
        self.search_box = QLineEdit()
        self.search_box.setObjectName("SearchBox")
        self.search_box.setPlaceholderText("🔎  Search anything...                                      Ctrl + K")
        self.search_box.returnPressed.connect(self.run_global_search)
        self.search_shortcut = QShortcut(QKeySequence("Ctrl+K"), self)
        self.search_shortcut.activated.connect(self.focus_search)
        topbar.addStretch(1)
        topbar.addWidget(self.search_box)
        topbar.addStretch(1)
        self.system_badge = QLabel("● Offline")
        self.system_badge.setObjectName("StatusBadge")
        topbar.addWidget(self.system_badge)
        layout.addLayout(topbar)

        self.pages = QStackedWidget()
        self.pages.setObjectName("PageStack")
        self.page_index = {}
        for name, builder in [
            ("Translate", self.build_translate_page),
            ("Reader", self.build_reader_page),
            ("Speech", self.build_speech_page),
            ("OCR", self.build_ocr_page),
            ("Notes", self.build_notes_page),
            ("Settings", self.build_settings_page),
        ]:
            self.page_index[name] = self.pages.addWidget(builder())
        layout.addWidget(self.pages, 1)
        return area

    def build_right_panel(self):
        panel = QWidget()
        panel.setObjectName("RightPanel")
        panel.setFixedWidth(350)
        self.right_layout = QVBoxLayout(panel)
        self.right_layout.setContentsMargins(12, 86, 26, 22)
        self.right_layout.setSpacing(16)
        return panel

    def build_footer(self):
        footer = QFrame()
        footer.setObjectName("Footer")
        footer.setFixedHeight(72)
        layout = QHBoxLayout(footer)
        layout.setContentsMargins(28, 0, 28, 0)
        self.footer_left = QLabel("● Ready")
        self.footer_left.setObjectName("FooterBadge")
        self.footer_center = QLabel("All processing is done offline on your device")
        self.footer_center.setObjectName("Muted")
        self.footer_right = QLabel("🛡 Offline Mode")
        self.footer_right.setObjectName("FooterText")
        layout.addWidget(self.footer_left)
        layout.addStretch(1)
        layout.addWidget(self.footer_center)
        layout.addStretch(1)
        layout.addWidget(self.footer_right)
        return footer

    # ---------- Styling ----------
    def apply_style(self):
        self.setStyleSheet("""
            #Root, QMainWindow { background: #0A1120; color: #F8FAFC; font-family: Segoe UI; font-size: 14px; }
            #Sidebar { background: #070D19; border-right: 1px solid rgba(148,163,184,0.12); }
            #MainArea, #RightPanel { background: #0A1120; }
            #Logo { background: qlineargradient(x1:0,y1:0,x2:1,y2:1, stop:0 #8B5CF6, stop:1 #06B6D4); border-radius: 13px; font-size: 26px; padding: 9px; }
            #AppTitle { font-size: 23px; font-weight: 800; color: #F8FAFC; }
            #Muted { color: #9CA3AF; font-size: 13px; }
            #SectionLabel { color: #94A3B8; font-size: 12px; font-weight: 700; letter-spacing: 1px; }
            #LinkLabel { color: #3B82F6; font-size: 12px; }
            #NavButton { text-align: left; background: transparent; color: #E5E7EB; border: none; border-radius: 14px; padding: 15px 18px; font-size: 15px; }
            #NavButton:hover { background: #111C2F; }
            #NavButton:checked { background: qlineargradient(x1:0,y1:0,x2:1,y2:0, stop:0 #2563EB, stop:1 #1E40AF); color: white; font-weight: 700; }
            #PrimaryButton { background: qlineargradient(x1:0,y1:0,x2:1,y2:0, stop:0 #2563EB, stop:1 #4F46E5); color: white; border: none; border-radius: 14px; padding: 13px 18px; font-weight: 700; }
            #PrimaryButton:hover { background: #3B82F6; }
            #SecondaryButton { background: #162238; color: #E5E7EB; border: 1px solid rgba(148,163,184,0.16); border-radius: 12px; padding: 10px 14px; }
            #SecondaryButton:hover { border: 1px solid #3B82F6; }
            #SearchBox { background: #0F172A; border: 1px solid rgba(148,163,184,0.20); border-radius: 15px; padding: 12px 18px; color: #E5E7EB; min-width: 430px; }
            #StatusBadge, #FooterBadge { background: rgba(34,197,94,0.12); color: #86EFAC; border-radius: 14px; padding: 9px 14px; }
            #Footer { background: #0B1424; border-top: 1px solid rgba(148,163,184,0.12); }
            #FooterText { color: #D1D5DB; }
            #PageTitle { font-size: 26px; font-weight: 800; color: #F8FAFC; }
            #Card { background: #121C2F; border: 1px solid rgba(148,163,184,0.16); border-radius: 20px; }
            #SmallCard { background: #121C2F; border: 1px solid rgba(148,163,184,0.14); border-radius: 16px; }
            #CardTitle { font-size: 13px; font-weight: 800; color: #E5E7EB; letter-spacing: .5px; }
            #BigText { font-size: 18px; font-weight: 700; color: #F8FAFC; }
            QTextEdit { background: #0F172A; color: #F8FAFC; border: 1px solid rgba(148,163,184,0.22); border-radius: 16px; padding: 14px; selection-background-color: #1E40AF; }
            QComboBox { background: #0F172A; color: #F8FAFC; border: 1px solid rgba(148,163,184,0.24); border-radius: 14px; padding: 9px 13px; min-height: 32px; }
            QComboBox:hover { border: 1px solid #3B82F6; }
            #SpeedBox { min-width: 92px; max-width: 110px; padding: 7px 10px; }
            QCheckBox { color: #D1D5DB; spacing: 8px; }
            QLabel { color: #E5E7EB; }
        """)

    # ---------- Shared helpers ----------
    def run_background(self, task, on_success):
        future = self.executor.submit(task)
        timer = QTimer(self)
        timer.setInterval(100)

        def check_future():
            if not future.done():
                return
            timer.stop()
            timer.deleteLater()
            try:
                on_success(future.result())
            except Exception as exc:
                self.set_status(f"Error: {exc}", error=True)
                if hasattr(self, "system_badge"):
                    self.system_badge.setText("● Backend Offline")

        timer.timeout.connect(check_future)
        timer.start()

    def set_status(self, message, error=False):
        self.footer_left.setText(("⚠ " if error else "● ") + message)

    def ensure_audio_backend(self) -> bool:
        if getattr(self, "audio_backend_ready", False):
            return True
        message = "Audio backend unavailable"
        if getattr(self, "audio_backend_error", None):
            message += f": {self.audio_backend_error}"
        self.set_status(message, error=True)
        return False

    def focus_search(self):
        self.search_box.setFocus()
        self.search_box.selectAll()
        self.set_status("Search focused")

    def run_global_search(self):
        query = self.search_box.text().strip().lower()
        if not query:
            self.set_status("Enter a search term")
            return

        matches = []
        for page, items in self.recent_history.items():
            for item in items:
                if query in item[1].lower():
                    matches.append((page, item[1]))

        if matches:
            page, text = matches[0]
            self.switch_page(page)
            self.set_status(f"Found in {page}: {text}")
        else:
            self.set_status(f"No recent match for: {query}")

    def add_recent(self, page_name, label, icon=None):
        if page_name not in self.recent_history:
            self.recent_history[page_name] = []
        default_icons = {
            "Translate": "🌐", "Reader": "📄", "OCR": "▣",
            "Speech": "🎙", "Notes": "📝", "Settings": "⚙",
        }
        entry = (icon or default_icons.get(page_name, "•"), str(label), "now")
        old_items = [item for item in self.recent_history[page_name] if item[1] != entry[1]]
        self.recent_history[page_name] = [entry] + old_items[:5]
        if self.current_page_name == page_name:
            self.update_recent_items(page_name)

    def language_box(self, default="en", include_auto=False):
        box = QComboBox()
        if include_auto:
            box.addItem("Auto", "auto")
        for label, code in LANGUAGES:
            box.addItem(label, code)
        idx = box.findData(default)
        if idx >= 0:
            box.setCurrentIndex(idx)
        return box

    def page_title(self, title, subtitle=None):
        wrap = QVBoxLayout()
        label = QLabel(title)
        label.setObjectName("PageTitle")
        wrap.addWidget(label)
        if subtitle:
            sub = QLabel(subtitle)
            sub.setObjectName("Muted")
            wrap.addWidget(sub)
        return wrap

    def switch_page(self, page_name):
        self.current_page_name = page_name
        self.pages.setCurrentIndex(self.page_index[page_name])
        for name, button in self.nav_buttons.items():
            button.setChecked(name == page_name)
        self.update_right_panel(page_name)
        self.update_recent_items(page_name)

    def clear_layout(self, layout):
        while layout.count():
            item = layout.takeAt(0)
            widget = item.widget()
            if widget:
                widget.deleteLater()
            child = item.layout()
            if child:
                self.clear_layout(child)

    def update_recent_items(self, page_name):
        if not hasattr(self, "recent_items_box"):
            return
        self.clear_layout(self.recent_items_box)
        items = self.recent_history.get(page_name, [])
        if not items:
            items = [("•", f"No recent {page_name.lower()} items", "")]
        for icon, text, time_label in items[:6]:
            row = QHBoxLayout()
            name = QLabel(f"{icon}  {text}")
            name.setObjectName("Muted")
            stamp = QLabel(time_label)
            stamp.setObjectName("Muted")
            row.addWidget(name)
            row.addStretch(1)
            row.addWidget(stamp)
            self.recent_items_box.addLayout(row)

    def info_card(self, title, rows):
        card = Card("SmallCard")
        layout = QVBoxLayout(card)
        layout.setContentsMargins(18, 16, 18, 16)
        layout.setSpacing(10)
        heading = QLabel(title)
        heading.setObjectName("CardTitle")
        layout.addWidget(heading)
        for left, right in rows:
            row = QHBoxLayout()
            l = QLabel(left)
            l.setObjectName("Muted")
            r = QLabel(str(right))
            r.setAlignment(Qt.AlignRight)
            row.addWidget(l)
            row.addStretch(1)
            row.addWidget(r)
            layout.addLayout(row)
        return card

    def update_right_panel(self, page_name):
        self.clear_layout(self.right_layout)
        if page_name == "Translate":
            self.right_layout.addWidget(self.info_card("QUICK ACTIONS", [("Copy output", "Ready"), ("Save note", "Ready")]))
            source = self.translate_source.currentText() if hasattr(self, "translate_source") else "Auto"
            target = self.translate_target.currentText() if hasattr(self, "translate_target") else "—"
            self.right_layout.addWidget(self.info_card("TRANSLATION INFO", [("Source", source), ("Target", target), ("Route", getattr(self, "last_translation_route", "—"))]))
        elif page_name == "Reader":
            self.right_layout.addWidget(self.info_card("READER INFO", [("File", self.reader_current_file), ("Language", self.reader_detected_language), ("Voice", self.reader_tts_lang.currentText() if hasattr(self, "reader_tts_lang") else self.reader_voice_label)]))
            speed = self.reader_speed.currentText() if hasattr(self, "reader_speed") else "1.0x"
            self.right_layout.addWidget(self.info_card("PLAYBACK", [("Speed", speed), ("Status", self.reader_playback_status), ("Duration", self.format_ms(self.reader_audio_duration_ms))]))
        elif page_name == "OCR":
            self.right_layout.addWidget(self.info_card("OCR INFO", [("Language", self.ocr_lang.currentText() if hasattr(self, "ocr_lang") else "Auto"), ("Input", "Image/PDF"), ("Status", getattr(self, "ocr_status", "Ready"))]))
        elif page_name == "Speech":
            self.right_layout.addWidget(self.info_card("SPEECH", [("Recorder", "Ready"), ("STT", "Offline"), ("Export", "Available")]))
        elif page_name == "Notes":
            self.right_layout.addWidget(self.info_card("NOTES", [("Storage", "SQLite"), ("Search", "Ctrl+K")]))
        else:
            self.right_layout.addWidget(self.info_card("SYSTEM", [("Backend", "Checking"), ("Mode", "Offline")]))
        self.right_layout.addStretch(1)

    # ---------- Translate ----------
    def build_translate_page(self):
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(14)
        layout.addLayout(self.page_title("Translate", "Translate text offline with language-aware views."))

        controls = QGridLayout()
        self.translate_source = self.language_box("auto", include_auto=True)
        self.translate_target = self.language_box("de")
        controls.addWidget(QLabel("Source"), 0, 0)
        controls.addWidget(self.translate_source, 0, 1)
        controls.addWidget(QLabel("Target"), 0, 2)
        controls.addWidget(self.translate_target, 0, 3)
        controls.setColumnStretch(1, 1)
        controls.setColumnStretch(3, 1)
        layout.addLayout(controls)

        self.translate_input = QTextEdit()
        self.translate_input.setPlaceholderText("Enter or paste text...")
        self.translate_output = QTextEdit()
        self.translate_output.setReadOnly(True)
        self.translate_output.setPlaceholderText("Translation output will appear here...")

        button_row = QHBoxLayout()
        import_doc_btn = QPushButton("＋ Import Document")
        import_doc_btn.setObjectName("SecondaryButton")
        import_doc_btn.clicked.connect(self.import_translate_document)

        btn = QPushButton("Translate")
        btn.setObjectName("PrimaryButton")
        btn.clicked.connect(self.translate_text)

        export_txt_btn = QPushButton("Export TXT")
        export_txt_btn.setObjectName("SecondaryButton")
        export_txt_btn.clicked.connect(lambda: self.export_translation("txt"))

        export_docx_btn = QPushButton("Export DOCX")
        export_docx_btn.setObjectName("SecondaryButton")
        export_docx_btn.clicked.connect(lambda: self.export_translation("docx"))

        export_pdf_btn = QPushButton("Export PDF")
        export_pdf_btn.setObjectName("SecondaryButton")
        export_pdf_btn.clicked.connect(lambda: self.export_translation("pdf"))

        button_row.addWidget(import_doc_btn)
        button_row.addWidget(btn)
        button_row.addStretch(1)
        button_row.addWidget(export_txt_btn)
        button_row.addWidget(export_docx_btn)
        button_row.addWidget(export_pdf_btn)

        layout.addWidget(QLabel("Input"))
        layout.addWidget(self.translate_input, 1)
        layout.addLayout(button_row)
        layout.addWidget(QLabel("Output"))
        layout.addWidget(self.translate_output, 1)

        translate_player = Card("Card")
        translate_player_layout = QVBoxLayout(translate_player)
        translate_player_layout.setContentsMargins(18, 14, 18, 14)

        translate_player_top = QHBoxLayout()
        self.translate_time_left = QLabel("00:00")
        self.translate_time_left.setObjectName("BigText")
        self.translate_time_right = QLabel("00:00")
        self.translate_time_right.setObjectName("BigText")
        self.translate_waveform = AudioWaveform()
        self.translate_waveform.seek_requested.connect(self.seek_reader_to_progress)
        translate_player_top.addWidget(self.translate_time_left)
        translate_player_top.addWidget(self.translate_waveform, 1)
        translate_player_top.addWidget(self.translate_time_right)

        translate_controls = QHBoxLayout()
        for label, handler in [
            ("⏪ Rewind", self.rewind_audio),
            ("▶ Read Translation", self.read_translation_aloud),
            ("⏸ Pause / Resume", self.pause_resume_audio),
            ("■ Stop", self.stop_audio),
        ]:
            control = QPushButton(label)
            control.setObjectName("SecondaryButton" if "Read" not in label else "PrimaryButton")
            control.clicked.connect(handler)
            translate_controls.addWidget(control)

        translate_controls.addStretch(1)
        translate_speed_group = QHBoxLayout()
        translate_speed_group.setSpacing(8)
        translate_speed_label = QLabel("Speed")
        translate_speed_label.setObjectName("Muted")
        self.translate_speed = QComboBox()
        self.translate_speed.setObjectName("SpeedBox")
        for label, value in [("0.75x", 0.75), ("1.0x", 1.0), ("1.25x", 1.25), ("1.5x", 1.5)]:
            self.translate_speed.addItem(label, value)
        self.translate_speed.setCurrentIndex(1)
        self.translate_speed.currentIndexChanged.connect(self.change_playback_speed)
        translate_speed_group.addWidget(translate_speed_label)
        translate_speed_group.addWidget(self.translate_speed)
        translate_controls.addLayout(translate_speed_group)

        translate_player_layout.addLayout(translate_player_top)
        translate_player_layout.addLayout(translate_controls)
        layout.addWidget(translate_player)
        return page

    def simplify_route(self, route):
        if not route:
            return "—"
        compact = []
        for item in route:
            if item and (not compact or compact[-1] != item):
                compact.append(item)
        if len(compact) > 2:
            return f"{compact[0]} → {compact[-1]}"
        return " → ".join(compact) or "—"

    def format_translation_output(self, translation):
        # Phase 3: keep normal output clean and export-ready. Auxiliary Hindi
        # romanization/devanagari views are intentionally hidden from the main
        # document workflow so they do not contaminate TXT/DOCX/PDF exports.
        return (translation.get("translated_text", "") or "").strip()

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            for url in event.mimeData().urls():
                suffix = Path(url.toLocalFile()).suffix.lower()
                if suffix in DOCUMENT_EXTENSIONS or suffix in AUDIO_EXTENSIONS or suffix in OCR_EXTENSIONS:
                    event.acceptProposedAction()
                    return
        event.ignore()

    def dropEvent(self, event):
        if not event.mimeData().hasUrls():
            event.ignore()
            return

        first = event.mimeData().urls()[0].toLocalFile()
        if not first:
            event.ignore()
            return

        suffix = Path(first).suffix.lower()
        if suffix in AUDIO_EXTENSIONS:
            self.import_speech_audio_path(first)
        elif suffix in IMAGE_EXTENSIONS or self.current_page_name == "OCR" and suffix == ".pdf":
            self.ocr_extract_path(first)
        elif self.current_page_name == "Reader":
            self.import_reader_document_path(first)
        else:
            self.import_translate_document_path(first)

        event.acceptProposedAction()

    def import_translate_document(self):
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Open document for translation",
            "",
            "Documents (*.txt *.md *.pdf *.docx *.rtf *.html *.htm *.csv *.json *.xml)",
        )
        if path:
            self.import_translate_document_path(path)

    def import_translate_document_path(self, path: str):
        self.switch_page("Translate")
        self.translate_output.setText("")
        self.translate_input.setText("Importing document...")
        self.translate_document_name = Path(path).name
        self.translate_document_source_path = str(Path(path).resolve())
        self.set_status(f"Importing {self.translate_document_name}")

        def task():
            with open(path, "rb") as file:
                response = requests.post(
                    f"{SERVER_URL}/reader/import",
                    files={"file": file},
                    data={"lang": self.translate_source.currentData()},
                    timeout=240,
                )
            return response.json()

        def success(data):
            if not data.get("ok"):
                self.translate_input.setText(str(data))
                self.set_status("Document import failed", error=True)
                return

            detected = data.get("detected_language", {}) or {}
            detected_lang = detected.get("language")
            # Keep Source on Auto for mixed-language imports. Otherwise the UI can
            # accidentally force the whole file through one stale route.
            if detected_lang and not detected.get("is_mixed"):
                idx = self.translate_source.findData(detected_lang)
                if idx >= 0:
                    self.translate_source.setCurrentIndex(idx)

            self.translate_input.setText(data.get("text", ""))
            self.add_recent("Translate", self.translate_document_name, "📄")
            self.set_status(f"Document ready: {self.translate_document_name}")
            self.update_right_panel("Translate")

        self.run_background(task, success)

    def can_use_format_preserving_export(self, file_type: str) -> bool:
        if file_type not in {"txt", "docx", "pdf"}:
            return False
        if not self.translate_document_source_path:
            return False
        return Path(self.translate_document_source_path).suffix.lower() in {".txt", ".md", ".csv", ".docx", ".pdf"}

    def export_format_preserving_translation(self, file_type: str, save_path: str):
        if not self.translate_document_source_path:
            raise RuntimeError("No imported source document available for format-preserving export.")

        with open(self.translate_document_source_path, "rb") as source_file:
            response = requests.post(
                f"{SERVER_URL}/translate/document/export",
                files={"file": source_file},
                data={
                    "source_lang": self.translate_source.currentData(),
                    "target_lang": self.translate_target.currentData(),
                    "output_format": file_type,
                },
                timeout=600,
            )

        if response.status_code != 200:
            raise RuntimeError(response.text)

        content_type = response.headers.get("content-type", "")
        if "application/json" in content_type:
            payload = response.json()
            if not payload.get("ok", False):
                raise RuntimeError(payload.get("error") or str(payload))

        Path(save_path).write_bytes(response.content)

    def export_translation(self, file_type: str):
        text = (self.last_translation_plain_text or self.translate_output.toPlainText()).strip()
        if not text or text == "Translating...":
            self.set_status("No translation to export", error=True)
            return

        filters = {
            "txt": "Text File (*.txt)",
            "docx": "Word Document (*.docx)",
            "pdf": "PDF File (*.pdf)",
        }
        default_name = "translated_text." + file_type
        if self.translate_document_name:
            default_name = Path(self.translate_document_name).stem + "_translated." + file_type

        path, _ = QFileDialog.getSaveFileName(self, "Export translation", default_name, filters[file_type])
        if not path:
            return

        try:
            if self.can_use_format_preserving_export(file_type):
                self.set_status(f"Exporting {file_type.upper()} with source formatting")
                self.export_format_preserving_translation(file_type, path)
            elif file_type == "txt":
                Path(path).write_text(text, encoding="utf-8")
            elif file_type == "docx":
                from docx import Document
                document = Document()
                document.add_heading("LinguaFusion Translation", level=1)
                for paragraph in text.split("\n"):
                    document.add_paragraph(paragraph)
                document.save(path)
            elif file_type == "pdf":
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.units import cm
                from xml.sax.saxutils import escape

                doc = SimpleDocTemplate(
                    path,
                    pagesize=A4,
                    rightMargin=2 * cm,
                    leftMargin=2 * cm,
                    topMargin=2 * cm,
                    bottomMargin=2 * cm,
                )
                styles = getSampleStyleSheet()
                story = [Paragraph("LinguaFusion Translation", styles["Title"]), Spacer(1, 0.4 * cm)]
                for paragraph in text.split("\n"):
                    paragraph = paragraph.strip()
                    if paragraph:
                        story.append(Paragraph(escape(paragraph), styles["BodyText"]))
                        story.append(Spacer(1, 0.2 * cm))
                doc.build(story)
            self.set_status(f"Exported {file_type.upper()}: {Path(path).name}")
        except Exception as exc:
            self.set_status(f"Export failed: {exc}", error=True)

    def translate_text(self):
        self.translate_output.setText("Translating...")
        self.set_status("Translating")

        def task():
            response = requests.post(
                f"{SERVER_URL}/reader/translate",
                data={
                    "text": self.translate_input.toPlainText(),
                    "source_lang": self.translate_source.currentData(),
                    "target_lang": self.translate_target.currentData(),
                },
                timeout=600,
            )
            return response.json()

        def success(data):
            if not data.get("ok"):
                message = data.get("error") or data.get("stage") or "Translation failed."
                self.translate_output.setText(f"Translation failed: {message}")
                self.set_status("Translation failed", error=True)
                return
            translation = data.get("translation", {})
            if not translation.get("ok", False):
                partial = (translation.get("translated_text") or "").strip()
                message = translation.get("error") or "Translation quality check failed."
                self.translate_output.setText((partial + "\n\n" if partial else "") + f"Translation failed: {message}")
                self.set_status("Translation failed", error=True)
                return
            self.last_translation_route = self.simplify_route(translation.get("route", []))
            formatted = self.format_translation_output(translation)
            self.last_translation_plain_text = formatted
            self.last_translation_tts_text = translation.get("translated_text", "")
            self.last_translation_input_snapshot = self.translate_input.toPlainText()
            self.last_translation_source_snapshot = self.translate_source.currentData()
            self.last_translation_target_snapshot = self.translate_target.currentData()
            self.translate_output.setText(formatted)
            label = self.translate_document_name or f"{data.get('source_lang', self.translate_source.currentData())} → {data.get('target_lang', self.translate_target.currentData())}"
            self.add_recent("Translate", label)
            self.set_status("Translation complete")
            self.update_right_panel("Translate")

        self.run_background(task, success)

    # ---------- Reader ----------
    def build_reader_page(self):
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.setSpacing(14)
        layout.addLayout(self.page_title("Reader", "Import documents, track reading progress, and listen with language-aware offline TTS."))

        controls = QGridLayout()
        self.reader_lang = self.language_box("auto", include_auto=True)
        self.reader_tts_lang = self.language_box("auto", include_auto=True)
        controls.addWidget(QLabel("OCR language"), 0, 0)
        controls.addWidget(self.reader_lang, 0, 1)
        controls.addWidget(QLabel("TTS language"), 0, 2)
        controls.addWidget(self.reader_tts_lang, 0, 3)
        controls.setColumnStretch(1, 1)
        controls.setColumnStretch(3, 1)
        layout.addLayout(controls)

        import_btn = QPushButton("Import Document")
        import_btn.setObjectName("PrimaryButton")
        import_btn.clicked.connect(self.import_document)
        layout.addWidget(import_btn)

        self.reader_info = QLabel("No document loaded.")
        self.reader_info.setObjectName("Muted")
        self.reader_stats = QLabel("Statistics: no document loaded.")
        self.reader_stats.setObjectName("Muted")

        tools_row = QHBoxLayout()
        seek_cursor_btn = QPushButton("▶ From Cursor")
        seek_cursor_btn.setObjectName("SecondaryButton")
        seek_cursor_btn.clicked.connect(self.seek_reader_to_cursor)
        tools_row.addWidget(seek_cursor_btn)
        tools_row.addStretch(1)

        export_txt_btn = QPushButton("Export TXT")
        export_txt_btn.setObjectName("SecondaryButton")
        export_txt_btn.clicked.connect(lambda: self.export_reader_text("txt"))
        export_docx_btn = QPushButton("Export DOCX")
        export_docx_btn.setObjectName("SecondaryButton")
        export_docx_btn.clicked.connect(lambda: self.export_reader_text("docx"))
        export_pdf_btn = QPushButton("Export PDF")
        export_pdf_btn.setObjectName("SecondaryButton")
        export_pdf_btn.clicked.connect(lambda: self.export_reader_text("pdf"))
        for w in [export_txt_btn, export_docx_btn, export_pdf_btn]:
            tools_row.addWidget(w)

        self.reader_text = QTextEdit()
        self.reader_text.cursorPositionChanged.connect(self.update_reader_cursor_status)
        layout.addWidget(self.reader_info)
        layout.addWidget(self.reader_stats)
        layout.addLayout(tools_row)
        layout.addWidget(self.reader_text, 1)

        player = Card("Card")
        player_layout = QVBoxLayout(player)
        player_layout.setContentsMargins(18, 14, 18, 14)

        player_top = QHBoxLayout()
        self.reader_time_left = QLabel("00:00")
        self.reader_time_left.setObjectName("BigText")
        self.reader_time_right = QLabel("00:00")
        self.reader_time_right.setObjectName("BigText")
        self.reader_waveform = AudioWaveform()
        self.reader_waveform.seek_requested.connect(self.seek_reader_to_progress)
        player_top.addWidget(self.reader_time_left)
        player_top.addWidget(self.reader_waveform, 1)
        player_top.addWidget(self.reader_time_right)

        controls_row = QHBoxLayout()
        self.reader_speed = QComboBox()
        self.reader_speed.setObjectName("SpeedBox")
        for label, value in [("0.75x", 0.75), ("1.0x", 1.0), ("1.25x", 1.25), ("1.5x", 1.5)]:
            self.reader_speed.addItem(label, value)
        self.reader_speed.setCurrentIndex(1)
        self.reader_speed.currentIndexChanged.connect(self.change_playback_speed)

        for label, handler in [("⏪ Rewind", self.rewind_audio), ("▶ Read", self.reader_speak), ("⏸ Pause / Resume", self.pause_resume_audio), ("■ Stop", self.stop_audio), ("⇄ Translate", self.send_reader_to_translate)]:
            b = QPushButton(label)
            b.setObjectName("SecondaryButton" if "Read" not in label else "PrimaryButton")
            b.clicked.connect(handler)
            controls_row.addWidget(b)

        controls_row.addStretch(1)
        speed_group = QHBoxLayout()
        speed_group.setSpacing(8)
        speed_label = QLabel("Speed")
        speed_label.setObjectName("Muted")
        speed_group.addWidget(speed_label)
        speed_group.addWidget(self.reader_speed)
        controls_row.addLayout(speed_group)

        player_layout.addLayout(player_top)
        player_layout.addLayout(controls_row)
        layout.addWidget(player)
        return page

    def import_document(self):
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Open document",
            "",
            "Documents (*.txt *.md *.pdf *.docx *.rtf *.html *.htm *.csv *.json *.xml)",
        )
        if path:
            self.import_reader_document_path(path)

    def import_reader_document_path(self, path: str):
        self.switch_page("Reader")
        self.reader_info.setText("Importing document...")
        self.reader_text.setText("Importing document...")
        self.set_status("Importing document")

        def task():
            with open(path, "rb") as f:
                response = requests.post(
                    f"{SERVER_URL}/reader/import",
                    files={"file": f},
                    data={"lang": self.reader_lang.currentData()},
                    timeout=240,
                )
            return response.json()

        def success(data):
            if not data.get("ok"):
                self.reader_info.setText("Import failed.")
                self.reader_text.setText(str(data))
                self.set_status("Import failed", error=True)
                return

            detected = data.get("detected_language", {})
            detected_lang = detected.get("language")
            confidence = detected.get("confidence", 0)
            self.reader_current_file = Path(path).name
            self.reader_detected_language = detected_lang or "Auto"
            self.reader_info.setText(
                f"File: {data.get('file_type')} | "
                f"Method: {data.get('method')} | "
                f"Detected: {detected_lang} ({round(confidence, 2)})"
            )
            # Keep TTS language on Auto by default. Auto now routes English,
            # German, Spanish and Hindi segments to the right Piper voice.
            self.reader_text.setPlainText(data["text"])
            self.reader_analysis = data.get("analysis", {}) or {}
            self.reader_sentence_ranges = self.reader_analysis.get("sentence_ranges", []) or []
            self.reader_current_sentence_index = -1
            self.reader_playback_sentence_offset = 0
            self.reader_playback_sentence_count = len(self.reader_sentence_ranges or [])
            self.reader_bookmarks = []
            self.reader_cursor_audio_cache = {}
            self.reader_full_audio_cache_key = None
            self.reader_full_audio_path = None
            self.update_reader_stats_label()
            self.clear_reader_highlight()
            self.add_recent("Reader", self.reader_current_file)
            self.set_status(f"{detected_lang or 'Language'} detected")
            self.update_right_panel("Reader")

        self.run_background(task, success)


    def update_reader_stats_label(self):
        if not hasattr(self, "reader_stats"):
            return
        analysis = self.reader_analysis or {}
        if not analysis:
            self.reader_stats.setText("Statistics: no document loaded.")
            return
        terms = analysis.get("technical_terms") or []
        term_preview = ", ".join(terms[:5]) if terms else "none"
        spoken_label = analysis.get("estimated_speaking_label")
        if not spoken_label:
            seconds = int(round(float(analysis.get("estimated_speaking_minutes", 0) or 0) * 60))
            spoken_label = self.format_duration_words(seconds)
        self.reader_stats.setText(
            f"Statistics: {analysis.get('words', 0)} words · "
            f"{analysis.get('sentences', 0)} sentences · "
            f"{analysis.get('paragraphs', 0)} paragraphs · "
            f"{spoken_label} spoken · Terms: {term_preview}"
        )

    def update_reader_cursor_status(self):
        if self.current_page_name != "Reader" or not hasattr(self, "reader_text"):
            return
        pos = self.reader_text.textCursor().position()
        sentence = self.find_reader_sentence_for_position(pos)
        if sentence and self.reader_playback_status != "Playing":
            self.set_status(f"Cursor at sentence {sentence.get('index', 0) + 1}")

    def find_reader_sentence_for_position(self, position: int):
        ranges = self.reader_sentence_ranges or []
        if not ranges:
            return None
        position = max(0, int(position or 0))
        for sentence in ranges:
            if int(sentence.get("start", 0)) <= position <= int(sentence.get("end", 0)):
                return sentence
        # If the cursor is in whitespace between sentences, choose the next
        # sentence; otherwise choose the closest previous sentence.
        for sentence in ranges:
            if position < int(sentence.get("start", 0)):
                return sentence
        return ranges[-1]

    def clear_reader_highlight(self):
        if hasattr(self, "reader_text"):
            self.reader_text.setExtraSelections([])
        self.reader_current_sentence_index = -1

    def highlight_reader_sentence(self, sentence_index: int):
        if not hasattr(self, "reader_text") or sentence_index < 0:
            return
        if sentence_index == self.reader_current_sentence_index:
            return
        if sentence_index >= len(self.reader_sentence_ranges):
            return
        sentence = self.reader_sentence_ranges[sentence_index]
        cursor = QTextCursor(self.reader_text.document())
        cursor.setPosition(int(sentence.get("start", 0)))
        cursor.setPosition(int(sentence.get("end", 0)), QTextCursor.KeepAnchor)
        fmt = QTextCharFormat()
        fmt.setBackground(QColor(59, 130, 246, 70))
        selection = QTextEdit.ExtraSelection()
        selection.cursor = cursor
        selection.format = fmt
        self.reader_text.setExtraSelections([selection])
        self.reader_current_sentence_index = sentence_index

    def reader_sentence_weight(self, item) -> int:
        text = item.get("text", "") or ""
        words = int(item.get("words", 0) or 0)
        if words <= 0:
            words = max(1, len(re.findall(r"[\w\u0900-\u097F]+", text)))
        # Short headings/table labels are spoken quickly. Long mixed-language
        # sentences need their natural word weight.
        if len(text.strip()) <= 40 and words <= 5:
            words = max(1, int(words * 0.65))
        return max(1, words)

    def reader_sentence_timeline(self, duration_ms: int | None = None):
        ranges = self.reader_sentence_ranges or []
        duration = int(duration_ms if duration_ms is not None else (self.reader_audio_duration_ms or 0))
        if not ranges or duration <= 0:
            return []
        weights = [self.reader_sentence_weight(item) for item in ranges]
        total = max(1, sum(weights))
        timeline = []
        running = 0
        for idx, weight in enumerate(weights):
            start = int((running / total) * duration)
            running += weight
            end = int((running / total) * duration)
            timeline.append((idx, max(0, start), max(start + 1, min(duration, end))))
        return timeline

    def update_reader_sentence_highlight(self, progress: float):
        if not self.reader_highlight_enabled or not self.reader_sentence_ranges:
            return
        duration = max(self.reader_audio_duration_ms or 1, 1)
        original_elapsed = max(0, min(duration, int(max(0.0, min(1.0, progress)) * duration)))
        timeline = self.reader_sentence_timeline(duration)
        if not timeline:
            return

        offset = max(0, int(getattr(self, "reader_playback_sentence_offset", 0) or 0))
        # When From Cursor starts with a small safety lead-in, keep the selected
        # sentence highlighted while playback is entering that sentence. This
        # avoids the UI jumping to the previous/next sentence after a cursor seek.
        if offset > 0 and offset < len(timeline):
            selected_start = timeline[offset][1]
            if original_elapsed <= selected_start + 900:
                self.highlight_reader_sentence(offset)
                return

        # Time-based highlight is more robust after pause/resume and seek than
        # local progress over the remaining text. Use a small lookahead because
        # pygame position updates lag slightly behind audible playback.
        adjusted = original_elapsed + 350
        index = timeline[-1][0]
        for idx, start, end in timeline:
            if start <= adjusted < end:
                index = idx
                break
        self.highlight_reader_sentence(index)

    def reader_full_tts_cache_key(self, text: str) -> str:
        basis = "|".join([
            self.reader_tts_lang.currentData() or "auto",
            str(self.reader_speed.currentData() or 1.0),
            hashlib.sha1((text or "").encode("utf-8", errors="ignore")).hexdigest(),
        ])
        return hashlib.sha1(basis.encode("utf-8")).hexdigest()

    def estimate_sentence_start_ms(self, sentence_index: int, duration_ms: int | None = None, cursor_safe: bool = False) -> int:
        ranges = self.reader_sentence_ranges or []
        if not ranges:
            return 0
        sentence_index = max(0, min(int(sentence_index or 0), len(ranges) - 1))
        duration = int(duration_ms if duration_ms is not None else (self.reader_audio_duration_ms or 0))
        timeline = self.reader_sentence_timeline(duration)
        if not timeline:
            return 0
        start = timeline[sentence_index][1]
        if cursor_safe and sentence_index > 0:
            end = timeline[sentence_index][2]
            sentence_span = max(1, end - start)
            guard = min(2200, max(700, int(sentence_span * 0.20)))
            start = max(0, start - guard)
        return max(0, min(max(duration - 1, 0), int(start)))

    def load_audio_base_without_playing(self, audio_path: str):
        self.reader_base_audio_path = str(Path(audio_path).resolve())
        self.reader_current_audio_path = self.reader_base_audio_path
        self.reader_seek_base_original_ms = 0
        self.reader_audio_duration_ms = len(AudioSegment.from_file(self.reader_base_audio_path))
        self.reader_active_duration_ms = self.effective_duration_ms()
        self.reader_time_left.setText("00:00")
        self.reader_time_right.setText(self.format_ms(self.reader_audio_duration_ms))
        self.reader_waveform.load_audio(self.reader_base_audio_path)
        self.reader_waveform.set_progress(0)
        if hasattr(self, "translate_time_left"):
            self.translate_time_left.setText("00:00")
            self.translate_time_right.setText(self.format_ms(self.reader_audio_duration_ms))
            self.translate_waveform.load_audio(self.reader_base_audio_path)
            self.translate_waveform.set_progress(0)

    def seek_reader_to_cursor(self):
        if not hasattr(self, "reader_text"):
            return
        full_text = self.reader_text.toPlainText()
        if not full_text.strip():
            self.set_status("No document loaded", error=True)
            return
        clean_full_text = self.clean_reader_text_for_tts(full_text)
        if not clean_full_text:
            self.set_status("No readable text", error=True)
            return

        pos = self.reader_text.textCursor().position()
        sentence = self.find_reader_sentence_for_position(pos)
        sentence_index = int(sentence.get("index", 0)) if sentence else 0
        if sentence:
            self.highlight_reader_sentence(sentence_index)

        full_cache_key = self.reader_full_tts_cache_key(clean_full_text)
        cached_full = (
            self.reader_full_audio_cache_key == full_cache_key
            and self.reader_full_audio_path
            and Path(self.reader_full_audio_path).exists()
        )

        self.reader_playback_sentence_offset = sentence_index
        self.reader_playback_sentence_count = max(1, len(self.reader_sentence_ranges or []) - sentence_index)

        if cached_full:
            if str(Path(self.reader_base_audio_path or "").resolve()) != str(Path(self.reader_full_audio_path).resolve()):
                self.load_audio_base_without_playing(self.reader_full_audio_path)
            start_ms = self.estimate_sentence_start_ms(sentence_index, cursor_safe=True)
            self.set_status(f"Playing from cursor at {self.format_ms(start_ms)}")
            self.play_current_audio_from(start_ms)
            return

        # First From Cursor use generates the full document audio once, then
        # seeks into that full cached waveform. This avoids creating separate
        # cursor-only audio and makes later cursor/seek behavior stable.
        self.reader_speak_text(
            clean_full_text,
            sentence_offset=sentence_index,
            full_cache_key=full_cache_key,
            play_from_sentence_index=sentence_index,
        )

    def add_reader_bookmark(self):
        if not hasattr(self, "reader_text"):
            return
        pos = self.reader_text.textCursor().position()
        text = self.reader_text.toPlainText()
        if not text.strip():
            self.set_status("No document loaded", error=True)
            return
        sentence = self.find_reader_sentence_for_position(pos)
        snippet = (sentence.get("text") if sentence else text[max(0, pos-40):pos+80]).strip()
        snippet = re.sub(r"\s+", " ", snippet)[:70]
        self.reader_bookmark_counter += 1
        label = f"B{self.reader_bookmark_counter}: {snippet or 'Position ' + str(pos)}"
        item = {"label": label, "position": pos, "snippet": snippet}
        self.reader_bookmarks.append(item)
        if hasattr(self, "reader_bookmark_box"):
            self.reader_bookmark_box.addItem(label, item)
            self.reader_bookmark_box.setCurrentIndex(self.reader_bookmark_box.count() - 1)
        self.set_status("Bookmark added")

    def go_reader_bookmark(self):
        if not hasattr(self, "reader_bookmark_box") or not hasattr(self, "reader_text"):
            return
        item = self.reader_bookmark_box.currentData()
        if not item:
            self.set_status("No bookmark selected", error=True)
            return
        pos = int(item.get("position", 0))
        cursor = self.reader_text.textCursor()
        cursor.setPosition(max(0, min(pos, len(self.reader_text.toPlainText()))))
        self.reader_text.setTextCursor(cursor)
        sentence = self.find_reader_sentence_for_position(pos)
        if sentence:
            self.highlight_reader_sentence(int(sentence.get("index", 0)))
        self.set_status(f"Bookmark opened: {item.get('label', '')[:40]}. Press From Cursor to read from here.")

    def export_reader_text(self, file_type: str):
        text = self.reader_text.toPlainText().strip() if hasattr(self, "reader_text") else ""
        if not text:
            self.set_status("No reader text to export", error=True)
            return
        filters = {"txt": "Text File (*.txt)", "docx": "Word Document (*.docx)", "pdf": "PDF File (*.pdf)"}
        default_name = Path(self.reader_current_file).stem if self.reader_current_file and self.reader_current_file != "None" else "reader_export"
        path, _ = QFileDialog.getSaveFileName(self, "Export reader document", f"{default_name}.{file_type}", filters[file_type])
        if not path:
            return
        self.set_status(f"Exporting reader {file_type.upper()}")

        def task():
            response = requests.post(
                f"{SERVER_URL}/reader/export",
                data={"text": text, "output_format": file_type, "title": default_name},
                timeout=240,
            )
            if response.status_code != 200:
                raise RuntimeError(response.text)
            if "application/json" in response.headers.get("content-type", ""):
                payload = response.json()
                if not payload.get("ok", False):
                    raise RuntimeError(payload.get("error") or str(payload))
            Path(path).write_bytes(response.content)
            return path

        def success(saved_path):
            self.set_status(f"Exported reader {file_type.upper()}: {Path(saved_path).name}")

        self.run_background(task, success)

    def clean_reader_text_for_tts(self, text):
        lines = []
        for line in text.splitlines():
            cleaned = line.strip()
            if not cleaned or cleaned.startswith("--- Page") or cleaned.startswith("--- OCR Page"):
                continue
            lines.append(re.sub(r"\s+", " ", cleaned))
        return " ".join(lines).strip()

    def reader_speak_text(self, text: str, sentence_offset: int = 0, cache_key: str | None = None, full_cache_key: str | None = None, play_from_sentence_index: int | None = None):
        text = self.clean_reader_text_for_tts(text)
        if not text:
            self.set_status("No text to read", error=True)
            return
        self.active_audio_context = "Reader"
        self.reader_playback_sentence_offset = max(0, int(sentence_offset or 0))
        self.reader_playback_sentence_count = max(1, len(self.reader_sentence_ranges or []) - self.reader_playback_sentence_offset)
        self.set_status("Generating reader audio")

        def task():
            response = requests.post(
                f"{SERVER_URL}/reader/speak",
                data={"text": text[:7000], "lang": self.reader_tts_lang.currentData(), "speed": "1.0"},
                timeout=360,
            )
            if response.status_code != 200:
                raise RuntimeError(response.text)
            out_path = Path(f"reader_output_{uuid4().hex}.wav").resolve()
            with open(out_path, "wb") as f:
                f.write(response.content)
            return str(out_path)

        def success(path):
            self.generated_audio_files.append(path)
            if cache_key:
                self.reader_cursor_audio_cache[cache_key] = path
            if full_cache_key:
                self.reader_full_audio_cache_key = full_cache_key
                self.reader_full_audio_path = path
            if play_from_sentence_index is not None:
                self.load_audio_base_without_playing(path)
                start_ms = self.estimate_sentence_start_ms(int(play_from_sentence_index), self.reader_audio_duration_ms, cursor_safe=True)
                self.reader_playback_sentence_offset = max(0, int(play_from_sentence_index))
                self.reader_playback_sentence_count = max(1, len(self.reader_sentence_ranges or []) - self.reader_playback_sentence_offset)
                self.play_current_audio_from(start_ms)
            else:
                self.play_audio_file(path)
            self.reader_playback_status = "Playing"
            self.set_status("Playing")
            self.update_right_panel("Reader")

        self.run_background(task, success)

    def reader_speak(self):
        self.reader_playback_sentence_offset = 0
        self.reader_playback_sentence_count = len(self.reader_sentence_ranges or [])
        text = self.reader_text.toPlainText() if hasattr(self, "reader_text") else ""
        clean_text = self.clean_reader_text_for_tts(text)
        full_cache_key = self.reader_full_tts_cache_key(clean_text) if clean_text else None
        self.reader_speak_text(text, sentence_offset=0, full_cache_key=full_cache_key)

    def format_ms(self, ms):
        if not ms or ms < 0:
            return "00:00"
        total_seconds = int(ms / 1000)
        return f"{total_seconds // 60:02d}:{total_seconds % 60:02d}"

    def format_duration_words(self, seconds):
        try:
            total = max(0, int(round(float(seconds))))
        except Exception:
            total = 0
        minutes, secs = divmod(total, 60)
        if minutes and secs:
            return f"{minutes} min {secs} sec"
        if minutes:
            return f"{minutes} min"
        return f"{secs} sec"

    def playback_rate(self) -> float:
        if getattr(self, "active_audio_context", "Reader") == "Translate" and hasattr(self, "translate_speed"):
            return float(self.translate_speed.currentData() or 1.0)
        if hasattr(self, "reader_speed"):
            return float(self.reader_speed.currentData() or 1.0)
        return 1.0

    def original_position_ms(self) -> int:
        """Return approximate position in the original, unmodified audio timeline."""
        if not self.reader_current_audio_path:
            return 0

        if not self.ensure_audio_backend():
            return self.reader_seek_base_original_ms

        if self.audio_paused:
            active_pos = self.reader_paused_pos_ms
        else:
            active_pos = max(pygame.mixer.music.get_pos(), 0)

        original_elapsed = self.reader_seek_base_original_ms + int(active_pos * max(self.reader_last_rate, 0.01))
        return max(0, min(original_elapsed, self.reader_audio_duration_ms or original_elapsed))

    def effective_duration_ms(self) -> int:
        if not self.reader_audio_duration_ms:
            return 0
        return int(self.reader_audio_duration_ms / max(self.playback_rate(), 0.01))

    def make_playback_segment_audio(self, source_path: str, speed: float, start_original_ms: int = 0) -> str:
        """
        Build the exact audio chunk pygame should play.

        pygame cannot reliably seek inside generated WAV files with play(start=...).
        So for seeking/dragging we create a temporary WAV segment starting at the
        requested original-audio position, then optionally apply ffmpeg atempo for
        speed control. This makes click/drag seek stable instead of jumping back to 0.
        """
        source = Path(source_path)
        speed = max(0.5, min(float(speed or 1.0), 2.0))
        start_original_ms = max(0, min(int(start_original_ms or 0), self.reader_audio_duration_ms or 0))

        if speed == 1.0 and start_original_ms == 0:
            return str(source)

        segment = AudioSegment.from_file(str(source))
        if start_original_ms > 0:
            segment = segment[start_original_ms:]

        temp_segment_path = Path(f"reader_segment_{uuid4().hex}.wav").resolve()
        segment.export(str(temp_segment_path), format="wav")
        self.generated_audio_files.append(str(temp_segment_path))

        if speed == 1.0:
            return str(temp_segment_path)

        adjusted_path = Path(f"reader_speed_{speed}_{uuid4().hex}.wav").resolve()

        subprocess.run(
            [
                "ffmpeg",
                "-y",
                "-i",
                str(temp_segment_path),
                "-filter:a",
                f"atempo={speed}",
                str(adjusted_path),
            ],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            check=True,
       )

        self.generated_audio_files.append(str(adjusted_path))
        return str(adjusted_path)

    def play_current_audio_from(self, original_position_ms: int = 0):
        if not self.ensure_audio_backend():
            return
        if not self.reader_base_audio_path:
            self.set_status("No audio loaded", error=True)
            return

        rate = self.playback_rate()
        self.reader_last_rate = rate
        self.reader_seek_base_original_ms = max(0, min(int(original_position_ms or 0), self.reader_audio_duration_ms or 0))
        self.reader_active_audio_path = self.make_playback_segment_audio(
            self.reader_base_audio_path,
            rate,
            self.reader_seek_base_original_ms,
        )
        self.reader_current_audio_path = self.reader_active_audio_path
        remaining_original = max(0, (self.reader_audio_duration_ms or 0) - self.reader_seek_base_original_ms)
        self.reader_active_duration_ms = int(remaining_original / max(rate, 0.01))

        pygame.mixer.music.stop()
        pygame.mixer.music.load(self.reader_active_audio_path)
        pygame.mixer.music.play()

        self.audio_paused = False
        self.reader_paused_pos_ms = 0
        self.reader_playback_status = "Playing"
        self.playback_timer.start()
        self.update_playback_ui()
        self.update_right_panel("Reader")

    def change_playback_speed(self):
        rate = self.playback_rate()
        if self.reader_base_audio_path:
            original_pos = self.original_position_ms()
            was_playing = self.ensure_audio_backend() and pygame.mixer.music.get_busy() and not self.audio_paused
            self.play_current_audio_from(original_pos)
            if not was_playing:
                pygame.mixer.music.pause()
                self.audio_paused = True
                self.reader_playback_status = "Paused"
        self.set_status(f"Playback speed set to {rate:g}x")
        self.update_playback_ui()
        if self.current_page_name == "Reader":
            self.update_right_panel("Reader")

    def update_playback_ui(self):
        if not self.reader_current_audio_path:
            return
        if not self.ensure_audio_backend():
            self.playback_timer.stop()
            return

        active_pos = self.reader_paused_pos_ms if self.audio_paused else max(pygame.mixer.music.get_pos(), 0)
        original_elapsed = self.reader_seek_base_original_ms + int(active_pos * max(self.reader_last_rate, 0.01))
        original_elapsed = max(0, min(original_elapsed, self.reader_audio_duration_ms or original_elapsed))

        duration = max(self.reader_audio_duration_ms or 1, 1)
        progress = max(0.0, min(1.0, original_elapsed / duration))

        self.reader_time_left.setText(self.format_ms(original_elapsed))
        self.reader_time_right.setText(self.format_ms(self.reader_audio_duration_ms))
        self.reader_waveform.set_progress(progress)
        if self.current_page_name == "Reader":
            self.update_reader_sentence_highlight(progress)

        if hasattr(self, "translate_time_left"):
            self.translate_time_left.setText(self.format_ms(original_elapsed))
            self.translate_time_right.setText(self.format_ms(self.reader_audio_duration_ms))
            self.translate_waveform.set_progress(progress)

        if progress >= 0.999 and not pygame.mixer.music.get_busy() and self.reader_playback_status == "Playing":
            self.reader_playback_status = "Stopped"
            self.audio_paused = False
            self.reader_paused_pos_ms = 0
            self.reader_seek_base_original_ms = 0
            self.playback_timer.stop()
            self.set_status("Audio finished")

        # Do not rebuild the right-side Reader cards on every playback tick.
        # Recreating those widgets at 60 FPS caused visible flicker while audio played.
        # The cards are refreshed only on discrete state changes: play, pause, stop, speed, import.

    def play_audio_file(self, audio_path):
        self.reader_base_audio_path = str(Path(audio_path).resolve())
        self.reader_current_audio_path = self.reader_base_audio_path
        self.reader_seek_base_original_ms = 0
        self.reader_audio_duration_ms = len(AudioSegment.from_file(self.reader_base_audio_path))
        self.reader_active_duration_ms = self.effective_duration_ms()
        self.reader_time_left.setText("00:00")
        self.reader_time_right.setText(self.format_ms(self.reader_audio_duration_ms))
        self.reader_waveform.load_audio(self.reader_base_audio_path)
        self.reader_waveform.set_progress(0)
        if hasattr(self, "translate_time_left"):
            self.translate_time_left.setText("00:00")
            self.translate_time_right.setText(self.format_ms(self.reader_audio_duration_ms))
            self.translate_waveform.load_audio(self.reader_base_audio_path)
            self.translate_waveform.set_progress(0)
        self.play_current_audio_from(0)

    def seek_reader_to_progress(self, progress: float):
        if not self.reader_base_audio_path or not self.reader_audio_duration_ms:
            self.set_status("No audio loaded", error=True)
            return

        progress = max(0.0, min(1.0, progress))
        original_position = int(self.reader_audio_duration_ms * progress)
        was_paused = self.audio_paused or self.reader_playback_status == "Paused"

        self.play_current_audio_from(original_position)

        if was_paused:
            pygame.mixer.music.pause()
            self.audio_paused = True
            self.reader_paused_pos_ms = 0
            self.reader_playback_status = "Paused"

        self.reader_waveform.set_progress(progress)
        self.reader_time_left.setText(self.format_ms(original_position))
        self.set_status(f"Jumped to {self.format_ms(original_position)}")
        self.update_right_panel("Reader")

    def rewind_audio(self):
        if not self.reader_base_audio_path:
            self.set_status("No audio loaded", error=True)
            return
        self.play_current_audio_from(0)
        self.reader_waveform.set_progress(0)
        self.set_status("Audio rewound")
        self.update_right_panel("Reader")

    def pause_resume_audio(self):
        if not self.ensure_audio_backend():
            return
        if not self.reader_current_audio_path:
            self.set_status("No audio loaded", error=True)
            return

        if self.audio_paused:
            pygame.mixer.music.unpause()
            self.audio_paused = False
            self.reader_playback_status = "Playing"
            self.playback_timer.start()
            self.set_status("Audio resumed")
        else:
            self.reader_paused_pos_ms = max(pygame.mixer.music.get_pos(), 0)
            pygame.mixer.music.pause()
            self.audio_paused = True
            self.reader_playback_status = "Paused"
            self.set_status("Audio paused")
        self.update_right_panel("Reader")

    def stop_audio(self):
        if self.ensure_audio_backend():
            pygame.mixer.music.stop()
        self.audio_paused = False
        self.reader_paused_pos_ms = 0
        self.reader_seek_base_original_ms = 0
        self.reader_playback_status = "Stopped"
        self.playback_timer.stop()
        if hasattr(self, "reader_time_left"):
            self.reader_time_left.setText("00:00")
            self.reader_waveform.set_progress(0)
            self.clear_reader_highlight()
        if hasattr(self, "translate_time_left"):
            self.translate_time_left.setText("00:00")
            self.translate_waveform.set_progress(0)
        self.set_status("Audio stopped")
        self.update_right_panel("Reader")

    def send_reader_to_translate(self):
        text = self.reader_text.toPlainText().strip()
        if not text:
            self.set_status("No reader text to send", error=True)
            return
        self.translate_input.setPlainText(text)
        detected = self.reader_detected_language if self.reader_detected_language and self.reader_detected_language != "Auto" else "auto"
        idx = self.translate_source.findData(detected)
        if idx >= 0:
            self.translate_source.setCurrentIndex(idx)
        self.switch_page("Translate")
        self.set_status("Reader text sent to Translate")

    def read_translation_aloud(self):
        current_input = self.translate_input.toPlainText().strip()
        current_source = self.translate_source.currentData()
        current_target = self.translate_target.currentData()

        if not current_input:
            self.set_status("No text to translate/read", error=True)
            return

        needs_translation = (
            not self.last_translation_tts_text.strip()
            or self.last_translation_input_snapshot != self.translate_input.toPlainText()
            or self.last_translation_source_snapshot != current_source
            or self.last_translation_target_snapshot != current_target
        )

        if needs_translation:
            self.translate_output.setText("Translating before reading...")
            self.set_status("Translating before reading")

            def translate_task():
                response = requests.post(
                    f"{SERVER_URL}/reader/translate",
                    data={
                        "text": self.translate_input.toPlainText(),
                        "source_lang": current_source,
                        "target_lang": current_target,
                    },
                    timeout=180,
                )
                return response.json()

            def translate_success(data):
                if not data.get("ok"):
                    self.translate_output.setText(str(data))
                    self.set_status("Translation failed", error=True)
                    return

                translation = data["translation"]
                self.last_translation_route = self.simplify_route(translation.get("route", []))
                formatted = self.format_translation_output(translation)
                self.last_translation_plain_text = formatted
                self.last_translation_tts_text = translation.get("translated_text", "")
                self.last_translation_input_snapshot = self.translate_input.toPlainText()
                self.last_translation_source_snapshot = current_source
                self.last_translation_target_snapshot = current_target
                self.translate_output.setText(formatted)
                label = self.translate_document_name or f"{data.get('source_lang', current_source)} → {data.get('target_lang', current_target)}"
                self.add_recent("Translate", label)
                self.update_right_panel("Translate")
                self.generate_translation_audio()

            self.run_background(translate_task, translate_success)
            return

        self.generate_translation_audio()

    def generate_translation_audio(self):
        text = (self.last_translation_tts_text or self.last_translation_plain_text or self.translate_output.toPlainText()).strip()
        if not text or text in {"Translating...", "Translating before reading..."}:
            self.set_status("No translation to read", error=True)
            return

        target_lang = self.translate_target.currentData() or "en"
        self.active_audio_context = "Translate"
        self.set_status("Generating translation audio")

        def task():
            response = requests.post(
                f"{SERVER_URL}/reader/speak",
                data={"text": text[:5000], "lang": target_lang, "speed": "1.0"},
                timeout=240,
            )
            if response.status_code != 200:
                raise RuntimeError(response.text)
            out_path = Path(f"translation_output_{uuid4().hex}.wav").resolve()
            with open(out_path, "wb") as audio_file:
                audio_file.write(response.content)
            return str(out_path)

        def success(path):
            self.generated_audio_files.append(path)
            self.play_audio_file(path)
            self.reader_playback_status = "Playing"
            self.set_status("Reading translation")
            self.update_right_panel(self.current_page_name)

        self.run_background(task, success)

    # ---------- OCR ----------
    def build_ocr_page(self):
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.addLayout(self.page_title("OCR", "Extract text from images and scanned documents."))
        self.ocr_lang = self.language_box("auto", include_auto=True)
        layout.addWidget(QLabel("OCR language"))
        layout.addWidget(self.ocr_lang)
        btn = QPushButton("Extract Text from Image")
        btn.setObjectName("PrimaryButton")
        btn.clicked.connect(self.ocr_extract)
        self.ocr_text = QTextEdit()
        layout.addWidget(btn)
        layout.addWidget(self.ocr_text, 1)
        return page

    def ocr_extract(self):
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Open image or scanned PDF",
            "",
            "Images/PDF (*.png *.jpg *.jpeg *.bmp *.tif *.tiff *.webp *.pdf);;All Files (*.*)",
        )
        if path:
            self.ocr_extract_path(path)

    def ocr_extract_path(self, path: str):
        if not path:
            return
        suffix = Path(path).suffix.lower()
        if suffix not in OCR_EXTENSIONS:
            self.set_status(f"Unsupported OCR file: {suffix}", error=True)
            return
        self.switch_page("OCR")
        self.ocr_text.setText("Extracting OCR text...")
        self.set_status("Running OCR")

        def task():
            with open(path, "rb") as f:
                response = requests.post(
                    f"{SERVER_URL}/ocr/extract",
                    files={"file": f},
                    data={"lang": self.ocr_lang.currentData()},
                    timeout=300,
                )
            return response.json()

        def success(data):
            details = []
            if data.get("average_confidence") is not None:
                details.append(f"Average confidence: {data.get('average_confidence')}%")
            if data.get("line_count") is not None:
                details.append(f"Lines: {data.get('line_count')}")
            body = data.get("text", str(data))
            if details and data.get("ok"):
                body = body + "\n\n--- OCR Diagnostics ---\n" + "\n".join(details)
            self.ocr_text.setPlainText(body)
            self.ocr_status = "Complete" if data.get("ok") else "Failed"
            if data.get("ok"):
                self.add_recent("OCR", Path(path).name)
            self.set_status("OCR complete" if data.get("ok") else "OCR failed", error=not data.get("ok"))
            self.update_right_panel("OCR")

        self.run_background(task, success)

    # ---------- Notes ----------
    def build_notes_page(self):
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.addLayout(self.page_title("Notes", "Save useful translations and extracted text."))
        self.note_title = QTextEdit()
        self.note_title.setMaximumHeight(52)
        self.note_content = QTextEdit()
        self.notes_list = QTextEdit()
        self.notes_list.setReadOnly(True)
        save_btn = QPushButton("Save Note")
        save_btn.setObjectName("PrimaryButton")
        save_btn.clicked.connect(self.save_note)
        refresh_btn = QPushButton("Refresh Notes")
        refresh_btn.setObjectName("SecondaryButton")
        refresh_btn.clicked.connect(self.load_notes)
        layout.addWidget(QLabel("Title"))
        layout.addWidget(self.note_title)
        layout.addWidget(QLabel("Content"))
        layout.addWidget(self.note_content, 1)
        layout.addWidget(save_btn)
        layout.addWidget(refresh_btn)
        layout.addWidget(QLabel("Saved Notes"))
        layout.addWidget(self.notes_list, 1)
        return page

    def save_note(self):
        def task():
            response = requests.post(f"{SERVER_URL}/notes/create", data={"title": self.note_title.toPlainText().strip() or "Untitled", "content": self.note_content.toPlainText(), "language": "en"}, timeout=30)
            return response.json()
        def success(data):
            title = data.get('title', 'Untitled')
            self.add_recent("Notes", title)
            self.set_status(f"Note saved: {title}")
        self.run_background(task, success)

    def load_notes(self):
        def task():
            return requests.get(f"{SERVER_URL}/notes", timeout=30).json()
        def success(notes):
            self.notes_list.setText("\n\n".join(f"[{n['id']}] {n['title']} ({n['language']})\n{n['content'][:300]}" for n in notes))
            self.set_status("Notes loaded")
        self.run_background(task, success)

    # ---------- Speech / Settings ----------
    def build_speech_page(self):
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.setSpacing(14)
        layout.addLayout(self.page_title("Speech", "Record audio with cleaner offline transcription, correction memory, and live translation."))

        controls = QGridLayout()
        self.speech_language = QComboBox()
        self.speech_language.addItem("Auto", "auto")
        for label, code in LANGUAGES:
            self.speech_language.addItem(label, code)
        self.speech_target_language = self.language_box("de")
        self.speech_live_translation = QCheckBox("Live Translation")
        self.speech_live_translation.setObjectName("Muted")
        self.speech_live_translation.setChecked(True)
        self.speech_smart_correction = QCheckBox("Smart Correction")
        self.speech_smart_correction.setObjectName("Muted")
        self.speech_smart_correction.setChecked(True)
        self.speech_correction_mode = QComboBox()
        self.speech_correction_mode.addItem("Offline only", "offline")
        self.speech_correction_mode.addItem("Free Auto", "free_auto")
        self.speech_correction_mode.addItem("LanguageTool", "languagetool")
        self.speech_correction_mode.addItem("Gemini", "gemini")
        self.speech_correction_mode.addItem("Groq", "groq")
        self.speech_correction_mode.addItem("OpenRouter", "openrouter")
        controls.addWidget(QLabel("STT language"), 0, 0)
        controls.addWidget(self.speech_language, 0, 1)
        controls.addWidget(QLabel("Translate to"), 0, 2)
        controls.addWidget(self.speech_target_language, 0, 3)
        controls.addWidget(self.speech_live_translation, 0, 4)
        controls.addWidget(self.speech_smart_correction, 1, 0)
        controls.addWidget(QLabel("Correction"), 1, 2)
        controls.addWidget(self.speech_correction_mode, 1, 3)
        controls.setColumnStretch(1, 1)
        controls.setColumnStretch(3, 1)
        layout.addLayout(controls)

        recorder = Card("Card")
        recorder_layout = QVBoxLayout(recorder)
        recorder_layout.setContentsMargins(18, 14, 18, 14)
        self.speech_status_label = QLabel("Ready to record or import audio.")
        self.speech_status_label.setObjectName("Muted")
        recorder_layout.addWidget(self.speech_status_label)

        wave_row = QHBoxLayout()
        self.speech_timer_label = QLabel("00:00")
        self.speech_timer_label.setObjectName("BigText")
        self.speech_waveform = AudioWaveform()
        wave_row.addWidget(self.speech_timer_label)
        wave_row.addWidget(self.speech_waveform, 1)
        recorder_layout.addLayout(wave_row)

        button_row = QHBoxLayout()
        record_btn = QPushButton("● Record")
        record_btn.setObjectName("PrimaryButton")
        record_btn.clicked.connect(self.start_speech_recording)
        pause_btn = QPushButton("⏸ Pause / Resume")
        pause_btn.setObjectName("SecondaryButton")
        pause_btn.clicked.connect(self.pause_resume_speech_recording)
        stop_btn = QPushButton("■ Stop")
        stop_btn.setObjectName("SecondaryButton")
        stop_btn.clicked.connect(self.stop_speech_recording)
        import_audio_btn = QPushButton("＋ Import Audio")
        import_audio_btn.setObjectName("SecondaryButton")
        import_audio_btn.clicked.connect(self.import_speech_audio)
        full_song_btn = QPushButton("🎵 Full Song Pass")
        full_song_btn.setObjectName("SecondaryButton")
        full_song_btn.clicked.connect(self.full_song_pass)
        reference_lyrics_btn = QPushButton("📄 Reference Lyrics")
        reference_lyrics_btn.setObjectName("SecondaryButton")
        reference_lyrics_btn.clicked.connect(self.reference_lyrics_pass)
        for btn in [record_btn, pause_btn, stop_btn, import_audio_btn, full_song_btn, reference_lyrics_btn]:
            button_row.addWidget(btn)
        recorder_layout.addLayout(button_row)
        layout.addWidget(recorder)

        action_row = QHBoxLayout()
        correct_btn = QPushButton("＋ Add Correction")
        correct_btn.setObjectName("SecondaryButton")
        correct_btn.clicked.connect(self.add_manual_correction)
        smart_fix_btn = QPushButton("✨ Smart Fix Selection")
        smart_fix_btn.setObjectName("SecondaryButton")
        smart_fix_btn.clicked.connect(self.smart_fix_selection)
        read_transcript_btn = QPushButton("▶ Read Transcript")
        read_transcript_btn.setObjectName("PrimaryButton")
        read_transcript_btn.clicked.connect(self.read_speech_transcript)
        read_translation_btn = QPushButton("▶ Read Translation")
        read_translation_btn.setObjectName("SecondaryButton")
        read_translation_btn.clicked.connect(self.read_speech_translation)
        pause_audio_btn = QPushButton("⏸ Pause / Resume")
        pause_audio_btn.setObjectName("SecondaryButton")
        pause_audio_btn.clicked.connect(self.pause_resume_audio)
        stop_audio_btn = QPushButton("■ Stop")
        stop_audio_btn.setObjectName("SecondaryButton")
        stop_audio_btn.clicked.connect(self.stop_audio)
        rewind_audio_btn = QPushButton("⏪ Restart")
        rewind_audio_btn.setObjectName("SecondaryButton")
        rewind_audio_btn.clicked.connect(self.rewind_audio)
        export_txt_btn = QPushButton("Export TXT")
        export_txt_btn.setObjectName("SecondaryButton")
        export_txt_btn.clicked.connect(lambda: self.export_speech_text("txt"))
        export_docx_btn = QPushButton("Export DOCX")
        export_docx_btn.setObjectName("SecondaryButton")
        export_docx_btn.clicked.connect(lambda: self.export_speech_text("docx"))
        export_pdf_btn = QPushButton("Export PDF")
        export_pdf_btn.setObjectName("SecondaryButton")
        export_pdf_btn.clicked.connect(lambda: self.export_speech_text("pdf"))
        save_note_btn = QPushButton("Save Note")
        save_note_btn.setObjectName("SecondaryButton")
        save_note_btn.clicked.connect(self.save_speech_as_note)
        for btn in [correct_btn, smart_fix_btn, read_transcript_btn, read_translation_btn, pause_audio_btn, stop_audio_btn, rewind_audio_btn]:
            action_row.addWidget(btn)
        action_row.addStretch(1)
        for btn in [export_txt_btn, export_docx_btn, export_pdf_btn, save_note_btn]:
            action_row.addWidget(btn)
        layout.addLayout(action_row)

        layout.addWidget(QLabel("Transcript"))
        self.speech_transcript = QTextEdit()
        self.speech_transcript.setPlaceholderText("Live transcript appears here automatically while recording...")
        layout.addWidget(self.speech_transcript, 1)

        layout.addWidget(QLabel("Translation"))
        self.speech_translation = QTextEdit()
        self.speech_translation.setReadOnly(True)
        self.speech_translation.setPlaceholderText("Translated transcript will appear here...")
        layout.addWidget(self.speech_translation, 1)
        return page


    def load_corrections_from_backend(self):
        def task():
            try:
                return requests.get(f"{SERVER_URL}/corrections", timeout=10).json()
            except Exception:
                return {"ok": False, "corrections": {}}

        def success(data):
            if data.get("ok"):
                self.corrections = data.get("corrections", {}) or {}

        self.run_background(task, success)

    def load_smart_provider_status(self):
        def task():
            try:
                return requests.get(f"{SERVER_URL}/corrections/providers", timeout=10).json()
            except Exception as exc:
                return {"ok": False, "error": str(exc), "providers": {}}

        def success(data):
            self.smart_provider_status = data.get("providers", {}) if data.get("ok") else {}

        self.run_background(task, success)

    def current_smart_mode(self, for_live: bool = False) -> str:
        if for_live:
            # Keep live snapshots cheap and fast. Online semantic correction runs on final/import or selection.
            return "offline"
        if hasattr(self, "speech_smart_correction") and self.speech_smart_correction.isChecked():
            return self.speech_correction_mode.currentData() or self.ai_provider_config.get("default_mode", "free_auto")
        return "offline"

    def selected_text_widget(self):
        widget = QApplication.focusWidget()
        if isinstance(widget, QTextEdit) and widget.textCursor().hasSelection():
            return widget
        for candidate in [
            getattr(self, "speech_transcript", None),
            getattr(self, "speech_translation", None),
            getattr(self, "reader_text", None),
            getattr(self, "translate_output", None),
            getattr(self, "translate_input", None),
        ]:
            if isinstance(candidate, QTextEdit) and candidate.textCursor().hasSelection():
                return candidate
        return None

    def apply_local_corrections(self, text: str) -> str:
        corrected = text or ""
        for wrong, correct in sorted((self.corrections or {}).items(), key=lambda item: len(item[0]), reverse=True):
            if not wrong or not correct:
                continue
            corrected = re.sub(rf"(?<!\w){re.escape(wrong)}(?!\w)", correct, corrected, flags=re.IGNORECASE)
        return corrected.strip()

    def add_manual_correction(self):
        selected_widget = self.selected_text_widget()
        selected = selected_widget.textCursor().selectedText().replace("\u2029", " ").strip() if selected_widget else ""
        wrong, ok = QInputDialog.getText(self, "Add Correction", "Wrong text:", text=selected)
        if not ok or not wrong.strip():
            return
        correct, ok = QInputDialog.getText(self, "Add Correction", "Correct text:")
        if not ok or not correct.strip():
            return

        def task():
            response = requests.post(
                f"{SERVER_URL}/corrections/add",
                data={"wrong": wrong.strip(), "correct": correct.strip()},
                timeout=20,
            )
            return response.json()

        def success(data):
            if not data.get("ok"):
                self.set_status(f"Correction failed: {data.get('error')}", error=True)
                return
            self.corrections = data.get("corrections", {}) or {}
            transcript = self.apply_local_corrections(self.speech_transcript.toPlainText())
            self.speech_transcript.setText(transcript)
            self.last_speech_transcript = transcript
            self.set_status(f"Learned correction: {wrong.strip()} → {correct.strip()}")
            QMessageBox.information(self, "Correction saved", f"LinguaFusion will now replace:\n\n{wrong.strip()} → {correct.strip()}")

        self.run_background(task, success)

    def smart_fix_selection(self):
        widget = self.selected_text_widget()
        if not widget:
            self.set_status("Select text first, then use Smart Fix Selection", error=True)
            return

        selected = widget.textCursor().selectedText().replace("\u2029", " ").strip()
        if not selected:
            self.set_status("No selected text", error=True)
            return

        mode = self.speech_correction_mode.currentData() if hasattr(self, "speech_correction_mode") else "free_auto"
        if mode == "offline":
            mode = "free_auto"

        if mode not in {"offline", "languagetool"}:
            confirm = QMessageBox.question(
                self,
                "Online correction",
                f"Send this selected text to the selected free online provider?\n\n{selected[:500]}",
            )
            if confirm != QMessageBox.Yes:
                return

        self.set_status(f"Smart correcting selection with {mode}")

        def task():
            response = requests.post(
                f"{SERVER_URL}/corrections/smart",
                data={"text": selected, "mode": mode, "language": self.speech_language.currentData()},
                timeout=80,
            )
            return response.json()

        def success(data):
            if not data.get("ok"):
                self.set_status(f"Smart correction failed: {data.get('error')}", error=True)
                return

            corrected = (data.get("text") or selected).strip()
            provider = data.get("provider", mode)
            if corrected == selected:
                self.set_status(f"No change suggested by {provider}")
                return

            accept = QMessageBox.question(
                self,
                "Apply smart correction?",
                f"Provider: {provider}\n\nBefore:\n{selected}\n\nAfter:\n{corrected}\n\nApply this change?",
            )
            if accept != QMessageBox.Yes:
                self.set_status("Smart correction ignored")
                return

            cursor = widget.textCursor()
            cursor.insertText(corrected)
            widget.setTextCursor(cursor)

            learn = QMessageBox.question(
                self,
                "Learn correction?",
                "Save this selected correction to offline memory for future transcripts?",
            )
            if learn == QMessageBox.Yes:
                self.save_learned_correction(selected, corrected)
            else:
                self.set_status(f"Applied smart correction from {provider}")

        self.run_background(task, success)

    def save_learned_correction(self, wrong: str, correct: str):
        def task():
            response = requests.post(
                f"{SERVER_URL}/corrections/add",
                data={"wrong": wrong.strip(), "correct": correct.strip()},
                timeout=20,
            )
            return response.json()

        def success(data):
            if data.get("ok"):
                self.corrections = data.get("corrections", {}) or {}
                self.set_status("Smart correction learned offline")
            else:
                self.set_status(f"Learning failed: {data.get('error')}", error=True)

        self.run_background(task, success)

    def maybe_live_translate_speech(self, transcript: str):
        if not hasattr(self, "speech_live_translation") or not self.speech_live_translation.isChecked():
            return
        if not transcript or transcript in {"Transcribing...", "Listening... live transcription will update every few seconds."}:
            return

        def task():
            response = requests.post(
                f"{SERVER_URL}/reader/translate",
                data={
                    "text": transcript,
                    "source_lang": self.speech_language.currentData(),
                    "target_lang": self.speech_target_language.currentData(),
                },
                timeout=180,
            )
            return response.json()

        def success(data):
            if not data.get("ok"):
                return
            translation = data["translation"]
            formatted = self.format_translation_output(translation)
            self.last_speech_translation = translation.get("translated_text", "")
            self.speech_translation.setText(formatted)

        self.run_background(task, success)

    def write_speech_frames_to_temp_wav(self, frames, prefix: str = "speech_snapshot") -> str:
        if not frames:
            raise RuntimeError("No audio frames available.")
        audio = np.concatenate(frames, axis=0)
        out_path = Path(f"{prefix}_{uuid4().hex}.wav").resolve()
        sf.write(str(out_path), audio, self.speech_sample_rate)
        self.generated_audio_files.append(str(out_path))
        return str(out_path)

    def live_transcribe_speech_snapshot(self):
        if not self.speech_is_recording or self.speech_is_paused or self.speech_live_transcribing:
            return
        if len(self.speech_frames) < 25:
            return
        if len(self.speech_frames) == self.speech_last_live_frame_count:
            return

        frames_snapshot = [frame.copy() for frame in self.speech_frames]
        self.speech_last_live_frame_count = len(frames_snapshot)
        self.speech_live_transcribing = True
        self.speech_status_label.setText("Recording... live transcribing in background.")

        def task():
            wav_path = self.write_speech_frames_to_temp_wav(frames_snapshot, "speech_live")
            with open(wav_path, "rb") as audio_file:
                response = requests.post(
                    f"{SERVER_URL}/stt/transcribe",
                    files={"file": audio_file},
                    data={"language": self.speech_language.currentData(), "smart_mode": self.current_smart_mode(for_live=True)},
                    timeout=600,
                )
            return response.json()

        def success(data):
            self.speech_live_transcribing = False
            if not data.get("ok"):
                self.speech_status_label.setText("Recording... live transcription not available yet.")
                return
            transcript = self.apply_local_corrections(data.get("text", "").strip())
            if transcript:
                self.last_speech_transcript = transcript
                self.speech_transcript.setText(transcript)
                self.maybe_live_translate_speech(transcript)
                self.speech_status_label.setText("Recording... live transcript updating.")

        self.run_background(task, success)

    def update_speech_recording_ui(self):
        if self.speech_is_recording and not self.speech_is_paused and self.speech_recording_started_at:
            elapsed = self.speech_elapsed_when_paused + (time.time() - self.speech_recording_started_at)
        else:
            elapsed = self.speech_elapsed_when_paused
        total_seconds = int(max(0, elapsed))
        if hasattr(self, "speech_timer_label"):
            self.speech_timer_label.setText(f"{total_seconds // 60:02d}:{total_seconds % 60:02d}")
        if hasattr(self, "speech_waveform"):
            levels = self.speech_levels[-160:]
            if levels:
                maximum = max(levels) or 1.0
                self.speech_waveform.peaks = [min(1.0, value / maximum) for value in levels]
                self.speech_waveform.set_progress(1.0 if self.speech_is_recording else 0.0)

    def start_speech_recording(self):
        if self.speech_is_recording:
            self.set_status("Recording already active")
            return
        self.stop_audio()
        self.speech_frames = []
        self.speech_levels = []
        self.speech_current_wav_path = None
        self.speech_live_transcribing = False
        self.speech_last_live_frame_count = 0
        self.last_speech_transcript = ""
        self.last_speech_translation = ""
        self.last_speech_language = "auto"
        self.speech_is_recording = True
        self.speech_is_paused = False
        self.speech_elapsed_when_paused = 0.0
        self.speech_recording_started_at = time.time()
        if hasattr(self, "speech_transcript"):
            self.speech_transcript.setText("Listening... live transcription will update every few seconds.")
        if hasattr(self, "speech_translation"):
            self.speech_translation.setText("")

        def callback(indata, frames, time_info, status):
            if status:
                pass
            if self.speech_is_recording and not self.speech_is_paused:
                block = indata.copy()
                self.speech_frames.append(block)
                try:
                    self.speech_levels.append(float(np.sqrt(np.mean(block ** 2))))
                except Exception:
                    self.speech_levels.append(0.0)

        try:
            self.speech_stream = sd.InputStream(
                samplerate=self.speech_sample_rate,
                channels=self.speech_channels,
                dtype="float32",
                callback=callback,
            )
            self.speech_stream.start()
            self.speech_timer.start()
            self.speech_live_timer.start()
            self.speech_status_label.setText("Recording... speak now. Live transcription is automatic.")
            self.set_status("Recording")
        except Exception as exc:
            self.speech_is_recording = False
            self.speech_is_paused = False
            self.speech_status_label.setText("Microphone error.")
            self.set_status(f"Microphone error: {exc}", error=True)

    def pause_resume_speech_recording(self):
        if not self.speech_is_recording:
            self.set_status("No active recording to pause", error=True)
            return
        if self.speech_is_paused:
            self.speech_is_paused = False
            self.speech_recording_started_at = time.time()
            self.speech_live_timer.start()
            self.speech_status_label.setText("Recording resumed... live transcription active.")
            self.set_status("Recording resumed")
        else:
            if self.speech_recording_started_at:
                self.speech_elapsed_when_paused += time.time() - self.speech_recording_started_at
            self.speech_recording_started_at = None
            self.speech_is_paused = True
            self.speech_live_timer.stop()
            self.speech_status_label.setText("Recording paused. Press Pause / Resume to continue.")
            self.set_status("Recording paused")
        self.update_speech_recording_ui()

    def stop_speech_recording(self):
        if not self.speech_is_recording:
            self.set_status("No active recording")
            return
        if not self.speech_is_paused and self.speech_recording_started_at:
            self.speech_elapsed_when_paused += time.time() - self.speech_recording_started_at
        self.speech_recording_started_at = None
        self.speech_is_recording = False
        self.speech_is_paused = False
        self.speech_timer.stop()
        self.speech_live_timer.stop()
        try:
            if self.speech_stream:
                self.speech_stream.stop()
                self.speech_stream.close()
        except Exception:
            pass
        self.speech_stream = None

        if not self.speech_frames:
            self.speech_status_label.setText("No audio captured.")
            self.set_status("No audio captured", error=True)
            return

        out_path = self.write_speech_frames_to_temp_wav(self.speech_frames, "speech_recording")
        self.speech_current_wav_path = out_path
        self.speech_status_label.setText(f"Recording saved: {Path(out_path).name}. Final transcription running...")
        self.speech_waveform.load_audio(out_path)
        self.set_status("Recording stopped; transcribing automatically")
        self.transcribe_speech_audio(auto=True)

    def import_speech_audio(self):
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Open audio file",
            "",
            "Audio Files (*.wav *.mp3 *.m4a *.ogg *.flac *.aac *.wma);;All Files (*.*)",
        )
        if path:
            self.import_speech_audio_path(path)

    def import_speech_audio_path(self, path: str):
        if not path:
            return
        suffix = Path(path).suffix.lower()
        if suffix not in AUDIO_EXTENSIONS:
            self.set_status(f"Unsupported audio file: {suffix}", error=True)
            return
        self.switch_page("Speech")
        self.speech_current_wav_path = str(Path(path).resolve())
        self.speech_waveform.load_audio(self.speech_current_wav_path)
        try:
            duration = len(AudioSegment.from_file(self.speech_current_wav_path))
            self.speech_timer_label.setText(self.format_ms(duration))
        except Exception:
            pass
        self.speech_status_label.setText(f"Loaded audio: {Path(path).name}. Transcription running...")
        self.add_recent("Speech", Path(path).name, "🎙")
        self.set_status("Audio loaded by drag/drop; transcribing automatically")
        self.transcribe_speech_audio(auto=True)

    def full_song_pass(self):
        if self.speech_is_recording:
            self.stop_speech_recording()
            return

        if not self.speech_current_wav_path:
            self.set_status("No song/audio loaded for Full Song Pass", error=True)
            return

        self.speech_status_label.setText("Running Full Song Pass: full-file Whisper + provider comparison...")
        self.set_status("Full Song Pass running")
        self.speech_transcript.setText("Full Song Pass running... this may take longer than live STT.")

        def task():
            with open(self.speech_current_wav_path, "rb") as audio_file:
                response = requests.post(
                    f"{SERVER_URL}/stt/full-song-pass",
                    files={"file": audio_file},
                    data={
                        "language": self.speech_language.currentData(),
                        "smart_mode": self.speech_correction_mode.currentData() if hasattr(self, "speech_correction_mode") else "free_auto",
                    },
                    timeout=900,
                )
            return response.json()

        def success(data):
            if not data.get("ok"):
                self.speech_transcript.setText(str(data))
                self.set_status("Full Song Pass failed", error=True)
                return

            transcript = data.get("text", "").strip()
            self.last_speech_transcript = transcript
            self.last_speech_language = data.get("language", self.speech_language.currentData()) or "auto"
            self.speech_transcript.setText(transcript)
            provider = data.get("provider", "offline")
            candidates = data.get("candidates", []) or []
            engine = data.get("engine", "speech_engine_v2")
            best_score = "—"
            for candidate in candidates:
                if candidate.get("provider") == provider and candidate.get("score") is not None:
                    best_score = str(round(float(candidate.get("score", 0)), 1))
                    break
            self.speech_status_label.setText(
                f"Full Song Pass ready: selected {provider} from {len(candidates)} candidates · {engine} · score {best_score}"
            )
            self.set_status(f"Full Song Pass complete: {provider}")
            self.add_recent("Speech", "Full Song Pass", "🎵")

            if hasattr(self, "speech_live_translation") and self.speech_live_translation.isChecked():
                self.translate_speech_transcript()

        self.run_background(task, success)


    def reference_lyrics_pass(self):
        if self.speech_is_recording:
            self.stop_speech_recording()
            return

        if not self.speech_current_wav_path:
            self.set_status("No song/audio loaded for Reference Lyrics", error=True)
            return

        reference, ok = QInputDialog.getMultiLineText(
            self,
            "Reference Lyrics",
            "Paste the reference lyrics here. LinguaFusion will not fetch lyrics automatically.",
            "",
        )
        if not ok or not reference.strip():
            self.set_status("Reference lyrics cancelled")
            return

        self.speech_status_label.setText("Running Reference Lyrics Pass...")
        self.set_status("Reference Lyrics Pass running")
        self.speech_transcript.setText("Reference Lyrics Pass running...")

        def task():
            with open(self.speech_current_wav_path, "rb") as audio_file:
                response = requests.post(
                    f"{SERVER_URL}/stt/reference-lyrics-pass",
                    files={"file": audio_file},
                    data={
                        "reference_lyrics": reference,
                        "language": self.speech_language.currentData(),
                        "smart_mode": self.speech_correction_mode.currentData() if hasattr(self, "speech_correction_mode") else "free_auto",
                    },
                    timeout=900,
                )
            return response.json()

        def success(data):
            if not data.get("ok"):
                self.speech_transcript.setText(str(data))
                self.set_status("Reference Lyrics Pass failed", error=True)
                return

            transcript = data.get("text", "").strip()
            self.last_speech_transcript = transcript
            self.last_speech_language = data.get("language", "en") or "en"
            self.speech_transcript.setText(transcript)
            provider = data.get("provider", "reference_lyrics")
            accepted = data.get("reference_accepted", False)
            similarity = data.get("reference_similarity", "—")
            status = "accepted" if accepted else "not accepted"
            self.speech_status_label.setText(
                f"Reference Lyrics Pass ready: {status} · similarity {similarity} · {provider}"
            )
            self.set_status(f"Reference Lyrics Pass complete: {status}")
            self.add_recent("Speech", "Reference Lyrics Pass", "📄")

            if hasattr(self, "speech_live_translation") and self.speech_live_translation.isChecked():
                self.translate_speech_transcript()

        self.run_background(task, success)

    def transcribe_speech_audio(self, auto: bool = False):
        if self.speech_is_recording:
            self.stop_speech_recording()
            return
        if not self.speech_current_wav_path:
            self.set_status("No speech audio available", error=True)
            return
        if not auto or not self.last_speech_transcript:
            self.speech_transcript.setText("Transcribing...")
        self.set_status("Transcribing speech")

        def task():
            with open(self.speech_current_wav_path, "rb") as audio_file:
                response = requests.post(
                    f"{SERVER_URL}/stt/transcribe",
                    files={"file": audio_file},
                    data={"language": self.speech_language.currentData(), "smart_mode": self.current_smart_mode(for_live=False)},
                    timeout=600,
                )
            return response.json()

        def success(data):
            if not data.get("ok"):
                self.speech_transcript.setText(str(data))
                self.set_status("Transcription failed", error=True)
                return
            transcript = self.apply_local_corrections(data.get("text", "").strip())
            self.last_speech_transcript = transcript
            self.last_speech_language = data.get("language", self.speech_language.currentData()) or "auto"
            self.speech_transcript.setText(transcript)
            self.maybe_live_translate_speech(transcript)
            self.add_recent("Speech", "Transcript", "🎙")
            provider = data.get("provider", "offline")
            engine = data.get("engine", "speech_engine_v2")
            self.speech_status_label.setText(f"Final transcript ready: {self.last_speech_language} · {provider} · {engine}")
            self.set_status("Transcription complete")

        self.run_background(task, success)

    def translate_speech_transcript(self, auto_read: bool = False):
        transcript = self.speech_transcript.toPlainText().strip()
        if not transcript or transcript in {"Transcribing...", "Full Song Pass running... this may take longer than live STT."}:
            self.set_status("No transcript to translate", error=True)
            return

        # For music/full-song pass, the UI source may still be Auto. Use the detected language when available.
        source_lang = getattr(self, "last_speech_language", "auto") or self.speech_language.currentData()
        if source_lang == "auto":
            # Lyrics/songs tested so far are English; reader/translate works better with an explicit source.
            source_lang = "en"

        self.speech_translation.setText("Translating transcript...")
        self.set_status("Translating transcript")

        def task():
            response = requests.post(
                f"{SERVER_URL}/reader/translate",
                data={
                    "text": transcript,
                    "source_lang": source_lang,
                    "target_lang": self.speech_target_language.currentData(),
                },
                timeout=180,
            )
            return response.json()

        def success(data):
            translation = data.get("translation", {}) if isinstance(data, dict) else {}
            if not data.get("ok") or not translation.get("ok", False):
                error = translation.get("error") or data.get("error") or "Translation engine rejected this transcript."
                fallback = translation.get("translated_text", "") or ""
                message = f"Translation unavailable: {error}"
                if fallback:
                    message += "\n\nBest partial output:\n" + fallback
                self.speech_translation.setText(message)
                self.set_status("Transcript translation failed", error=True)
                return

            formatted = self.format_translation_output(translation)
            self.last_speech_translation = translation.get("translated_text", "")
            self.speech_translation.setText(formatted)
            self.set_status("Transcript translated")
            if auto_read:
                self.read_speech_translation()

        self.run_background(task, success)

    def generate_speech_audio_from_text(self, text: str, lang: str, status_label: str):
        if not text or text in {"Transcribing...", "Translating transcript..."}:
            self.set_status("No speech text to read", error=True)
            return
        if lang == "auto":
            lang = "en"
        self.active_audio_context = "Reader"
        self.set_status(status_label)

        def task():
            response = requests.post(
                f"{SERVER_URL}/reader/speak",
                data={"text": text[:5000], "lang": lang, "speed": "1.0"},
                timeout=240,
            )
            if response.status_code != 200:
                raise RuntimeError(response.text)
            out_path = Path(f"speech_tts_{uuid4().hex}.wav").resolve()
            with open(out_path, "wb") as audio_file:
                audio_file.write(response.content)
            return str(out_path)

        def success(path):
            self.generated_audio_files.append(path)
            self.play_audio_file(path)
            self.reader_playback_status = "Playing"
            self.set_status("Playing speech audio")

        self.run_background(task, success)

    def read_speech_transcript(self):
        text = self.speech_transcript.toPlainText().strip()
        self.generate_speech_audio_from_text(text, self.speech_language.currentData(), "Generating transcript audio")

    def read_speech_translation(self):
        text = (self.last_speech_translation or self.speech_translation.toPlainText()).strip()
        if not text or text == "Translating transcript...":
            self.translate_speech_transcript(auto_read=True)
            return
        self.generate_speech_audio_from_text(text, self.speech_target_language.currentData(), "Generating translation audio")

    def speech_export_text(self) -> str:
        transcript = self.speech_transcript.toPlainText().strip()
        translation = self.speech_translation.toPlainText().strip()
        parts = []
        if transcript and transcript != "Transcribing...":
            parts.append("Transcript\n----------\n" + transcript)
        if translation and translation != "Translating transcript...":
            parts.append("Translation\n----------\n" + translation)
        return "\n\n".join(parts).strip()

    def export_speech_text(self, file_type: str):
        text = self.speech_export_text()
        if not text:
            self.set_status("No speech text to export", error=True)
            return
        filters = {
            "txt": "Text File (*.txt)",
            "docx": "Word Document (*.docx)",
            "pdf": "PDF File (*.pdf)",
        }
        path, _ = QFileDialog.getSaveFileName(self, "Export speech text", f"speech_export.{file_type}", filters[file_type])
        if not path:
            return
        try:
            if file_type == "txt":
                Path(path).write_text(text, encoding="utf-8")
            elif file_type == "docx":
                from docx import Document
                doc = Document()
                doc.add_heading("LinguaFusion Speech Export", level=1)
                for paragraph in text.split("\n"):
                    doc.add_paragraph(paragraph)
                doc.save(path)
            elif file_type == "pdf":
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.units import cm
                from xml.sax.saxutils import escape
                doc = SimpleDocTemplate(path, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
                styles = getSampleStyleSheet()
                story = [Paragraph("LinguaFusion Speech Export", styles["Title"]), Spacer(1, 0.4*cm)]
                for paragraph in text.split("\n"):
                    if paragraph.strip():
                        story.append(Paragraph(escape(paragraph), styles["BodyText"]))
                        story.append(Spacer(1, 0.18*cm))
                doc.build(story)
            self.set_status(f"Speech export saved: {Path(path).name}")
        except Exception as exc:
            self.set_status(f"Export failed: {exc}", error=True)

    def save_speech_transcript(self):
        self.export_speech_text("txt")

    def save_speech_as_note(self):
        text = self.speech_export_text()
        if not text:
            self.set_status("No speech text to save", error=True)
            return
        def task():
            response = requests.post(
                f"{SERVER_URL}/notes/create",
                data={"title": "Speech note", "content": text, "language": self.speech_language.currentData()},
                timeout=30,
            )
            return response.json()
        def success(data):
            self.add_recent("Notes", data.get("title", "Speech note"), "📝")
            self.set_status("Speech saved as note")
        self.run_background(task, success)

    # ---------- Settings ----------
    def build_settings_page(self):
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.setSpacing(14)
        layout.addLayout(self.page_title("Settings", "Backend status, local configuration, and Smart Mode providers."))

        system_card = Card("Card")
        system_layout = QVBoxLayout(system_card)
        system_layout.setContentsMargins(18, 16, 18, 16)
        system_layout.addWidget(QLabel("System"))

        system_btns = QHBoxLayout()
        btn = QPushButton("Check Backend Health")
        btn.setObjectName("PrimaryButton")
        btn.clicked.connect(self.check_health)
        provider_btn = QPushButton("Check Smart Providers")
        provider_btn.setObjectName("SecondaryButton")
        provider_btn.clicked.connect(self.check_smart_providers)
        system_btns.addWidget(btn)
        system_btns.addWidget(provider_btn)
        system_btns.addStretch(1)
        system_layout.addLayout(system_btns)
        layout.addWidget(system_card)

        ai_card = Card("Card")
        ai_layout = QVBoxLayout(ai_card)
        ai_layout.setContentsMargins(18, 16, 18, 16)
        ai_layout.setSpacing(12)
        ai_title = QLabel("AI Providers / Smart Mode")
        ai_title.setObjectName("CardTitle")
        ai_layout.addWidget(ai_title)

        self.ai_smart_enabled = QCheckBox("Enable Smart Mode for speech final transcription and selected-text correction")
        self.ai_smart_enabled.setChecked(False)
        ai_layout.addWidget(self.ai_smart_enabled)

        mode_row = QGridLayout()
        self.ai_default_mode = QComboBox()
        self.ai_default_mode.addItem("Free Auto", "free_auto")
        self.ai_default_mode.addItem("LanguageTool only", "languagetool")
        self.ai_default_mode.addItem("Gemini", "gemini")
        self.ai_default_mode.addItem("Groq", "groq")
        self.ai_default_mode.addItem("OpenRouter", "openrouter")

        self.ai_primary_provider = QComboBox()
        self.ai_primary_provider.addItem("Gemini", "gemini")
        self.ai_primary_provider.addItem("Groq", "groq")
        self.ai_primary_provider.addItem("OpenRouter", "openrouter")

        mode_row.addWidget(QLabel("Correction mode"), 0, 0)
        mode_row.addWidget(self.ai_default_mode, 0, 1)
        mode_row.addWidget(QLabel("Primary provider"), 0, 2)
        mode_row.addWidget(self.ai_primary_provider, 0, 3)
        mode_row.setColumnStretch(1, 1)
        mode_row.setColumnStretch(3, 1)
        ai_layout.addLayout(mode_row)

        self.gemini_key_input = QLineEdit()
        self.gemini_key_input.setEchoMode(QLineEdit.Password)
        self.gemini_key_input.setPlaceholderText("Paste Gemini API key here")
        self.groq_key_input = QLineEdit()
        self.groq_key_input.setEchoMode(QLineEdit.Password)
        self.groq_key_input.setPlaceholderText("Paste Groq API key here")
        self.openrouter_key_input = QLineEdit()
        self.openrouter_key_input.setEchoMode(QLineEdit.Password)
        self.openrouter_key_input.setPlaceholderText("Paste OpenRouter API key here")

        key_grid = QGridLayout()
        key_grid.addWidget(QLabel("Gemini"), 0, 0)
        key_grid.addWidget(self.gemini_key_input, 0, 1)
        gemini_test = QPushButton("Test Gemini")
        gemini_test.setObjectName("SecondaryButton")
        gemini_test.clicked.connect(lambda: self.test_ai_provider("gemini"))
        key_grid.addWidget(gemini_test, 0, 2)

        key_grid.addWidget(QLabel("Groq"), 1, 0)
        key_grid.addWidget(self.groq_key_input, 1, 1)
        groq_test = QPushButton("Test Groq")
        groq_test.setObjectName("SecondaryButton")
        groq_test.clicked.connect(lambda: self.test_ai_provider("groq"))
        key_grid.addWidget(groq_test, 1, 2)

        key_grid.addWidget(QLabel("OpenRouter"), 2, 0)
        key_grid.addWidget(self.openrouter_key_input, 2, 1)
        openrouter_test = QPushButton("Test OpenRouter")
        openrouter_test.setObjectName("SecondaryButton")
        openrouter_test.clicked.connect(lambda: self.test_ai_provider("openrouter"))
        key_grid.addWidget(openrouter_test, 2, 2)
        key_grid.setColumnStretch(1, 1)
        ai_layout.addLayout(key_grid)

        save_row = QHBoxLayout()
        save_ai_btn = QPushButton("Save AI Provider Settings")
        save_ai_btn.setObjectName("PrimaryButton")
        save_ai_btn.clicked.connect(self.save_ai_provider_settings)
        load_ai_btn = QPushButton("Reload Settings")
        load_ai_btn.setObjectName("SecondaryButton")
        load_ai_btn.clicked.connect(self.load_ai_provider_settings)
        save_row.addWidget(save_ai_btn)
        save_row.addWidget(load_ai_btn)
        save_row.addStretch(1)
        ai_layout.addLayout(save_row)

        note = QLabel("Keys are stored locally in backend/storage/ai_providers.json. Do not commit/share this file.")
        note.setObjectName("Muted")
        ai_layout.addWidget(note)
        layout.addWidget(ai_card)

        self.health_output = QTextEdit()
        self.health_output.setReadOnly(True)
        layout.addWidget(self.health_output, 1)

        QTimer.singleShot(300, self.load_ai_provider_settings)
        return page

    def check_health(self):
        def task():
            response = requests.get(f"{SERVER_URL}/health", timeout=10)
            response.raise_for_status()
            return response.json()

        def success(data):
            services = data.get("services", {}) if isinstance(data, dict) else {}
            critical_ready = bool(services.get("speech")) and bool(services.get("tts"))
            self.system_badge.setText("● Offline Ready" if critical_ready else "● Backend Running")
            self.set_status("Ready" if critical_ready else "Backend running; check diagnostics")
            if hasattr(self, "health_output"):
                self.health_output.setText(json.dumps(data, indent=2, ensure_ascii=False))

        self.run_background(task, success)

    def check_smart_providers(self):
        def task():
            return requests.get(f"{SERVER_URL}/corrections/providers", timeout=10).json()

        def success(data):
            self.smart_provider_status = data.get("providers", {}) if data.get("ok") else {}
            if hasattr(self, "health_output"):
                self.health_output.setText(json.dumps(data, indent=2, ensure_ascii=False))
            self.set_status("Smart providers checked")

        self.run_background(task, success)

    def load_ai_provider_settings(self):
        if not hasattr(self, "ai_smart_enabled"):
            return

        def task():
            return requests.get(f"{SERVER_URL}/ai/providers/config", timeout=10).json()

        def success(data):
            if not data.get("ok"):
                self.set_status("Could not load AI provider settings", error=True)
                return
            self.ai_provider_config = data
            self.ai_smart_enabled.setChecked(bool(data.get("smart_mode_enabled")))
            idx = self.ai_default_mode.findData(data.get("default_mode", "free_auto"))
            if idx >= 0:
                self.ai_default_mode.setCurrentIndex(idx)
            idx = self.ai_primary_provider.findData(data.get("primary_provider", "gemini"))
            if idx >= 0:
                self.ai_primary_provider.setCurrentIndex(idx)

            has_keys = data.get("has_keys", {})
            self.gemini_key_input.setPlaceholderText("Gemini key saved" if has_keys.get("gemini") else "Paste Gemini API key here")
            self.groq_key_input.setPlaceholderText("Groq key saved" if has_keys.get("groq") else "Paste Groq API key here")
            self.openrouter_key_input.setPlaceholderText("OpenRouter key saved" if has_keys.get("openrouter") else "Paste OpenRouter API key here")

            if hasattr(self, "speech_smart_correction"):
                self.speech_smart_correction.setChecked(bool(data.get("smart_mode_enabled")))
            if hasattr(self, "speech_correction_mode"):
                idx = self.speech_correction_mode.findData(data.get("default_mode", "free_auto"))
                if idx >= 0:
                    self.speech_correction_mode.setCurrentIndex(idx)

            if hasattr(self, "health_output"):
                self.health_output.setText(json.dumps(data, indent=2, ensure_ascii=False))
            self.set_status("AI provider settings loaded")

        self.run_background(task, success)

    def save_ai_provider_settings(self):
        if not hasattr(self, "ai_smart_enabled"):
            return

        payload = {
            "smart_mode_enabled": str(self.ai_smart_enabled.isChecked()).lower(),
            "default_mode": self.ai_default_mode.currentData(),
            "primary_provider": self.ai_primary_provider.currentData(),
            "gemini_key": self.gemini_key_input.text().strip(),
            "groq_key": self.groq_key_input.text().strip(),
            "openrouter_key": self.openrouter_key_input.text().strip(),
        }

        def task():
            return requests.post(f"{SERVER_URL}/ai/providers/config", data=payload, timeout=20).json()

        def success(data):
            if not data.get("ok"):
                self.set_status("Could not save AI provider settings", error=True)
                return
            self.gemini_key_input.clear()
            self.groq_key_input.clear()
            self.openrouter_key_input.clear()
            self.load_ai_provider_settings()
            self.set_status("AI provider settings saved")

        self.run_background(task, success)

    def test_ai_provider(self, provider: str):
        self.save_ai_provider_settings()
        self.set_status(f"Testing {provider} provider")

        def delayed_test():
            time.sleep(0.5)
            return requests.post(f"{SERVER_URL}/ai/providers/test", data={"provider": provider}, timeout=90).json()

        def success(data):
            if hasattr(self, "health_output"):
                self.health_output.setText(json.dumps(data, indent=2, ensure_ascii=False))
            if data.get("ok"):
                self.set_status(f"{provider} test successful")
            else:
                self.set_status(f"{provider} test failed: {data.get('error')}", error=True)

        self.run_background(delayed_test, success)

    def closeEvent(self, event):
        try:
            self.playback_timer.stop()
            if hasattr(self, "speech_timer"):
                self.speech_timer.stop()
            if hasattr(self, "speech_live_timer"):
                self.speech_live_timer.stop()
            if getattr(self, "speech_stream", None):
                self.speech_stream.stop()
                self.speech_stream.close()
            if getattr(self, "audio_backend_ready", False):
                pygame.mixer.music.stop()
                pygame.mixer.quit()
        except Exception:
            pass
        for audio_file in self.generated_audio_files:
            try:
                os.remove(audio_file)
            except Exception:
                pass
        event.accept()


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = LinguaFusionWindow()
    window.show()
    sys.exit(app.exec())
